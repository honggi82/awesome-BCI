import csv
import html
import json
import math
import re
import shutil
import time
from collections import Counter, defaultdict
from datetime import date
from pathlib import Path

import requests
from docx import Document


ROOT = Path(__file__).resolve().parents[1]
DATA_DIR = ROOT / "data"
DOCS_DIR = ROOT / "docs"
PAPER_DIR = ROOT / "paper"
START_YEAR = 2000
END_YEAR = date.today().year
YEARS = list(range(START_YEAR, END_YEAR + 1))
YEAR_RANGE_TEXT = f"{START_YEAR}-{END_YEAR}"
YEAR_FILE_STEM = f"{START_YEAR}_{END_YEAR}"
PAPERS_JSON = f"papers_{YEAR_FILE_STEM}.json"
PAPERS_CSV = f"papers_{YEAR_FILE_STEM}.csv"
CANDIDATES_JSON = f"candidates_top500_{YEAR_FILE_STEM}.json"
CANDIDATES_CSV = f"candidates_top500_{YEAR_FILE_STEM}.csv"
TAXONOMY_CSV = f"papers_taxonomy_{YEAR_FILE_STEM}.csv"
PERIOD_ANALYSIS_JSON = f"period_analysis_{YEAR_FILE_STEM}.json"
TARGET_PER_YEAR = 100
CANDIDATES_PER_YEAR = 500

S2_BULK_URL = "https://api.semanticscholar.org/graph/v1/paper/search/bulk"
S2_FIELDS = ",".join([
    "paperId",
    "title",
    "year",
    "authors",
    "venue",
    "publicationVenue",
    "publicationDate",
    "citationCount",
    "influentialCitationCount",
    "abstract",
    "url",
    "externalIds",
    "openAccessPdf",
    "s2FieldsOfStudy",
    "publicationTypes",
])

QUERIES = [
    "brain computer interface",
    "brain machine interface",
    "BCI EEG",
    "motor imagery BCI",
    "SSVEP BCI",
    "P300 BCI",
    "brain computer interface rehabilitation",
    "invasive brain computer interface",
    "neural decoding brain computer interface",
    "deep learning brain computer interface",
]

RELEVANCE_PATTERNS = [
    r"\bbrain[- ]computer interface",
    r"\bbrain[- ]machine interface",
    r"\bBCI\b",
    r"\bSSVEP\b",
    r"\bP300\b",
    r"\bmotor imagery\b",
    r"\bneural decoding\b",
    r"\bintracortical\b",
    r"\belectrocorticography\b",
]

CATEGORIES = [
    ("Motor Imagery and Movement Decoding", ["motor imagery", "movement", "motor", "kinematic", "hand", "gait"]),
    ("SSVEP, P300, and ERP Spellers", ["ssvep", "p300", "erp", "speller", "steady state", "visual evoked"]),
    ("Rehabilitation and Neuroprosthetics", ["rehabilitation", "stroke", "prosthe", "exoskeleton", "restoration", "therapy"]),
    ("Invasive and Implantable Interfaces", ["invasive", "intracortical", "implant", "electrocorticography", "ecog", "microelectrode"]),
    ("Deep Learning and Representation Learning", ["deep learning", "convolution", "transformer", "graph neural", "domain adaptation", "self supervised"]),
    ("EEG Signal Processing and Datasets", ["eeg", "artifact", "signal processing", "dataset", "benchmark", "classification"]),
    ("Speech, Language, and Communication BCIs", ["speech", "language", "communication", "typing", "text", "decoding words"]),
    ("Hybrid, Affective, and Closed-loop BCIs", ["hybrid", "affective", "closed-loop", "closed loop", "neurofeedback", "adaptive"]),
]

TAXONOMY_VISUALS = {
    "Motor Imagery and Movement Decoding": {
        "motif": "motor",
        "accent": "#0f766e",
        "secondary": "#60a5fa",
    },
    "SSVEP, P300, and ERP Spellers": {
        "motif": "speller",
        "accent": "#7c3aed",
        "secondary": "#f59e0b",
    },
    "Rehabilitation and Neuroprosthetics": {
        "motif": "rehab",
        "accent": "#dc2626",
        "secondary": "#22c55e",
    },
    "Invasive and Implantable Interfaces": {
        "motif": "implant",
        "accent": "#2563eb",
        "secondary": "#a855f7",
    },
    "Deep Learning and Representation Learning": {
        "motif": "deep",
        "accent": "#9333ea",
        "secondary": "#14b8a6",
    },
    "EEG Signal Processing and Datasets": {
        "motif": "eeg",
        "accent": "#0891b2",
        "secondary": "#f97316",
    },
    "Speech, Language, and Communication BCIs": {
        "motif": "speech",
        "accent": "#be123c",
        "secondary": "#06b6d4",
    },
    "Hybrid, Affective, and Closed-loop BCIs": {
        "motif": "hybrid",
        "accent": "#16a34a",
        "secondary": "#8b5cf6",
    },
    "General BCI Methods and Systems": {
        "motif": "general",
        "accent": "#334155",
        "secondary": "#0f766e",
    },
}

KEYWORD_CONVENTION = [
    ("invasive", "Implanted or intracranial neural interfaces, including ECoG and intracortical recordings.", "2563eb"),
    ("non-invasive", "External sensing interfaces such as EEG, MEG, fNIRS, or fMRI.", "0f766e"),
    ("human", "Studies using human participants, patients, or volunteers.", "f59e0b"),
    ("non-human", "Animal, simulation, or non-human experimental settings.", "a855f7"),
    ("SMR", "Sensorimotor rhythm, ERD/ERS, or motor-imagery control paradigms.", "dc2626"),
    ("SSVEP", "Steady-state visual evoked potential paradigms.", "7c3aed"),
    ("P300", "P300 or event-related-potential speller paradigms.", "be123c"),
    ("arm-direction", "Arm, hand, reach, or directional movement decoding/control.", "0891b2"),
]
KEYWORD_COLORS = {keyword: color for keyword, _, color in KEYWORD_CONVENTION}

LANGUAGES = {
    "en": "English",
    "ko": "한국어",
    "zh": "中文",
    "ja": "日本語",
}

UI_LABELS = {
    "en": {
        "papers": "papers",
        "categories": "categories",
        "overview": "Category Overview",
        "limitations": "Limitations",
        "analysis": "Selected-period analysis",
        "totalSelected": "Total selected papers",
        "categoryCount": "Categories",
    },
    "ko": {
        "papers": "편",
        "categories": "개 분류",
        "overview": "분류 개요",
        "limitations": "한계",
        "analysis": "선택 기간 분석",
        "totalSelected": "총 선정 논문",
        "categoryCount": "분류",
    },
    "zh": {
        "papers": "篇论文",
        "categories": "个分类",
        "overview": "分类概览",
        "limitations": "局限性",
        "analysis": "所选期间分析",
        "totalSelected": "入选论文总数",
        "categoryCount": "分类数",
    },
    "ja": {
        "papers": "本",
        "categories": "分類",
        "overview": "カテゴリ概要",
        "limitations": "限界",
        "analysis": "選択期間の分析",
        "totalSelected": "選定論文総数",
        "categoryCount": "カテゴリ数",
    },
}

UI_LABELS["en"].update({
    "keyIdea": "Key idea",
    "strengths": "Strengths",
    "paperLimitations": "Limitations",
})
UI_LABELS["ko"].update({
    "keyIdea": "\ud575\uc2ec \uc544\uc774\ub514\uc5b4",
    "strengths": "\uc7a5\uc810",
    "paperLimitations": "\ud55c\uacc4",
})
UI_LABELS["zh"].update({
    "keyIdea": "\u6838\u5fc3\u601d\u60f3",
    "strengths": "\u4f18\u52bf",
    "paperLimitations": "\u5c40\u9650\u6027",
})
UI_LABELS["ja"].update({
    "keyIdea": "\u4e3b\u8981\u30a2\u30a4\u30c7\u30a2",
    "strengths": "\u5f37\u307f",
    "paperLimitations": "\u9650\u754c",
})

TAXONOMY_TRENDS = {
    "Motor Imagery and Movement Decoding": [
        "The field is moving from subject-specific pipelines toward cross-subject, calibration-light, and transfer-learning decoders for EEG motor imagery.",
        "Deep CNN, temporal convolution, graph, attention, and large EEG representation models are increasingly used to improve robustness under noisy and low-data conditions.",
        "Application work is expanding from binary hand imagery toward gait, lower-limb control, soft robotics, virtual feedback, and rehabilitation-oriented closed-loop use.",
    ],
    "SSVEP, P300, and ERP Spellers": [
        "Research is concentrating on high-speed, many-target communication systems with lower calibration burden and more stable target recognition.",
        "Training-free and adaptive spatial filtering, task-discriminant component analysis, and deep neural decoders are prominent directions for SSVEP/P300 reliability.",
        "Hybrid paradigms that combine SSVEP, P300, RSVP, EOG, or augmented/virtual reality interfaces are becoming a practical route to richer command sets.",
    ],
    "Rehabilitation and Neuroprosthetics": [
        "The dominant trend is integration of BCI with robotic gloves, exoskeletons, FES, VR, and task-oriented therapy for post-stroke and motor impairment rehabilitation.",
        "Studies increasingly ask whether BCI training transfers to activities of daily living rather than only improving offline decoding accuracy.",
        "Recent work points toward home-use, patient-centered protocols, multimodal feedback, and combined motor-cognitive-affective rehabilitation.",
    ],
    "Invasive and Implantable Interfaces": [
        "Invasive BCI research is shifting toward high-bandwidth, stable, long-term decoding for movement, communication, and sensory feedback.",
        "Key engineering themes include wireless operation, power efficiency, signal longevity, surgical risk, and reliability outside tightly controlled laboratory sessions.",
        "Clinical translation is increasingly tied to home use, user safety, tactile feedback, speech decoding, and realistic functional tasks.",
    ],
    "Deep Learning and Representation Learning": [
        "Deep learning work is moving beyond single-dataset CNN classifiers toward temporal, spectral, graph, transformer, and attention-based architectures.",
        "A major trend is representation learning that can transfer across users, sessions, headsets, and BCI paradigms with less subject-specific calibration.",
        "Interpretability, uncertainty, robustness to artifacts, and benchmark comparability are becoming as important as peak classification accuracy.",
    ],
    "EEG Signal Processing and Datasets": [
        "This taxonomy emphasizes reproducible preprocessing, artifact handling, channel selection, spatial filtering, and benchmark datasets for EEG-based BCI.",
        "The field is gradually shifting from isolated algorithm papers toward shared datasets, standardized evaluation, and metadata-aware comparisons.",
        "Hybrid EEG/fNIRS, transfer learning, and open benchmark resources are recurring themes for improving generalization and clinical relevance.",
    ],
    "Speech, Language, and Communication BCIs": [
        "Communication BCI is expanding from spelling paradigms toward imagined speech, decoded language, and higher-bandwidth text production.",
        "Both invasive and non-invasive studies are exploring more naturalistic communication, including speech motor cortex decoding and inner-speech EEG datasets.",
        "The central challenge remains preserving accuracy, latency, vocabulary size, and user autonomy in real-world assistive communication.",
    ],
    "Hybrid, Affective, and Closed-loop BCIs": [
        "Hybrid BCI combines multiple signals or paradigms to improve reliability, command diversity, and asynchronous control.",
        "Closed-loop and neurofeedback work is increasingly focused on user adaptation, mental-state awareness, fatigue, affect, and training protocols.",
        "The trend is toward systems that adapt to the user over time rather than treating decoding as a one-shot offline classification problem.",
    ],
    "General BCI Methods and Systems": [
        "General BCI work is consolidating definitions, system architectures, evaluation principles, and long-term challenges across invasive and non-invasive approaches.",
        "Recent surveys increasingly emphasize translation, usability, ethics, safety, reproducibility, and the gap between laboratory performance and real-world use.",
        "This area functions as the conceptual bridge between signal processing, neural engineering, clinical deployment, and human-centered design.",
    ],
}

TAXONOMY_LIMITATIONS = {
    "Motor Imagery and Movement Decoding": [
        "Cross-subject and cross-session variability still limits real-world robustness, especially when calibration time is short.",
        "Many high-scoring methods remain validated on offline datasets rather than sustained closed-loop control or clinical rehabilitation outcomes.",
        "Citation-ranked lists can favor mature EEG motor imagery pipelines over newer low-citation work on multimodal movement decoding.",
    ],
    "SSVEP, P300, and ERP Spellers": [
        "Visual fatigue, gaze dependence, and stimulus comfort remain practical barriers for long-duration communication use.",
        "High-speed results often depend on controlled displays, known target layouts, and calibration conditions that may not transfer to daily use.",
        "Hybrid paradigms improve command diversity but add setup complexity and make fair benchmarking harder.",
    ],
    "Rehabilitation and Neuroprosthetics": [
        "Clinical evidence is often fragmented across small cohorts, heterogeneous protocols, and short follow-up windows.",
        "Improvements in decoder accuracy or therapy-session metrics do not always demonstrate transfer to activities of daily living.",
        "Patient selection, therapist involvement, adverse-event reporting, and outcome measures vary enough to limit direct comparison.",
    ],
    "Invasive and Implantable Interfaces": [
        "Surgical risk, long-term signal stability, device maintenance, and participant burden remain major translation barriers.",
        "Many studies involve small cohorts or case reports, so headline performance can be difficult to generalize.",
        "Home deployment, cybersecurity, informed consent, explantation, and support infrastructure remain difficult to evaluate consistently.",
    ],
    "Deep Learning and Representation Learning": [
        "Performance can be inflated by dataset leakage, weak cross-subject splits, or inconsistent preprocessing across benchmarks.",
        "Large models often improve accuracy while reducing interpretability, uncertainty awareness, and clinical trust.",
        "Low-data and noisy-session robustness remains unresolved for many architectures outside curated datasets.",
    ],
    "EEG Signal Processing and Datasets": [
        "Benchmark datasets vary widely in task design, sensors, preprocessing, and participant populations.",
        "Artifact handling, channel selection, and evaluation protocols are not standardized enough for simple leaderboard-style comparison.",
        "Participant counts, trial structure, hardware, and licensing differences can limit reproducible reuse across laboratories.",
    ],
    "Speech, Language, and Communication BCIs": [
        "Vocabulary size, latency, privacy, and user autonomy remain difficult to balance in practical assistive communication.",
        "The strongest decoding results often rely on invasive recordings or carefully constrained tasks with limited participant diversity.",
        "Non-invasive imagined-speech datasets are still comparatively small and hard to evaluate consistently.",
    ],
    "Hybrid, Affective, and Closed-loop BCIs": [
        "Combining paradigms can improve reliability but increases calibration, hardware, and user workload.",
        "Closed-loop adaptation is hard to evaluate because user learning, fatigue, and affect change during use.",
        "Longitudinal real-world studies remain scarce, so durability and user acceptance are not well captured.",
    ],
    "General BCI Methods and Systems": [
        "Survey and system papers can dominate citation-ranked views while obscuring smaller empirical advances.",
        "Evaluation language remains inconsistent across paradigms, making taxonomy boundaries imperfect.",
        "Broad system claims often need stronger protocol-level reproducibility checks and real-world usability validation.",
    ],
}

CATEGORY_LOCALIZATION = {
    "Motor Imagery and Movement Decoding": {
        "en": {
            "name": "Motor Imagery and Movement Decoding",
            "focus": "movement intention decoding, motor imagery EEG, gait, hand control, and rehabilitation-oriented feedback",
            "challenge": "cross-subject variability and offline-only validation remain the main interpretation risks",
        },
        "ko": {
            "name": "운동상상 및 움직임 디코딩",
            "focus": "운동 의도 해석, 운동상상 EEG, 보행·손 제어, 재활 피드백",
            "challenge": "개인차와 세션 차이, 오프라인 검증 중심 결과가 핵심 해석 한계입니다",
        },
        "zh": {
            "name": "运动想象与运动解码",
            "focus": "运动意图解码、运动想象 EEG、步态和手部控制、康复反馈",
            "challenge": "跨被试差异、跨会话不稳定以及偏离线验证是主要解释风险",
        },
        "ja": {
            "name": "運動イメージと運動デコーディング",
            "focus": "運動意図の解読、運動イメージ EEG、歩行・手の制御、リハビリ向けフィードバック",
            "challenge": "被験者差、セッション差、オフライン評価中心である点が主な解釈上の制約です",
        },
    },
    "SSVEP, P300, and ERP Spellers": {
        "en": {
            "name": "SSVEP, P300, and ERP Spellers",
            "focus": "visual evoked responses, ERP spelling, high-speed target selection, and calibration-light communication",
            "challenge": "visual fatigue, gaze dependence, and controlled display assumptions limit daily-use transfer",
        },
        "ko": {
            "name": "SSVEP, P300, ERP 스펠러",
            "focus": "시각유발반응, ERP 스펠링, 고속 표적 선택, 저보정 의사소통",
            "challenge": "시각 피로, 시선 의존성, 통제된 디스플레이 조건이 실사용 전이를 제한합니다",
        },
        "zh": {
            "name": "SSVEP、P300 与 ERP 拼写器",
            "focus": "视觉诱发反应、ERP 拼写、高速目标选择、低校准通信",
            "challenge": "视觉疲劳、凝视依赖和受控显示条件限制了日常使用迁移",
        },
        "ja": {
            "name": "SSVEP・P300・ERP スペラー",
            "focus": "視覚誘発反応、ERP スペリング、高速ターゲット選択、低キャリブレーション通信",
            "challenge": "視覚疲労、視線依存、制御された表示環境への依存が日常利用への移行を制約します",
        },
    },
    "Rehabilitation and Neuroprosthetics": {
        "en": {
            "name": "Rehabilitation and Neuroprosthetics",
            "focus": "post-stroke therapy, robotic assistance, FES, exoskeletons, neuroprosthetic control, and functional recovery",
            "challenge": "small cohorts and heterogeneous protocols make clinical effect size difficult to compare",
        },
        "ko": {
            "name": "재활 및 신경보철",
            "focus": "뇌졸중 후 재활, 로봇 보조, FES, 외골격, 신경보철 제어, 기능 회복",
            "challenge": "소규모 코호트와 이질적 프로토콜 때문에 임상 효과 비교가 어렵습니다",
        },
        "zh": {
            "name": "康复与神经假体",
            "focus": "卒中后治疗、机器人辅助、FES、外骨骼、神经假体控制和功能恢复",
            "challenge": "小样本和异质方案使临床效果大小难以比较",
        },
        "ja": {
            "name": "リハビリテーションと神経補綴",
            "focus": "脳卒中後リハビリ、ロボット支援、FES、外骨格、神経補綴制御、機能回復",
            "challenge": "小規模コホートと不均一なプロトコルにより臨床効果の比較が難しくなります",
        },
    },
    "Invasive and Implantable Interfaces": {
        "en": {
            "name": "Invasive and Implantable Interfaces",
            "focus": "intracortical arrays, ECoG, implantable devices, high-bandwidth decoding, sensory feedback, and home use",
            "challenge": "surgical risk, signal longevity, device maintenance, and participant burden remain central constraints",
        },
        "ko": {
            "name": "침습형 및 이식형 인터페이스",
            "focus": "피질내 전극, ECoG, 이식형 장치, 고대역폭 디코딩, 감각 피드백, 가정 사용",
            "challenge": "수술 위험, 장기 신호 안정성, 장치 유지관리, 참여자 부담이 핵심 제약입니다",
        },
        "zh": {
            "name": "侵入式与植入式接口",
            "focus": "皮层内阵列、ECoG、植入设备、高带宽解码、感觉反馈和家庭使用",
            "challenge": "手术风险、信号寿命、设备维护和参与者负担仍是核心约束",
        },
        "ja": {
            "name": "侵襲型・植込み型インターフェース",
            "focus": "皮質内アレイ、ECoG、植込みデバイス、高帯域デコーディング、感覚フィードバック、在宅利用",
            "challenge": "手術リスク、信号の長期安定性、デバイス保守、参加者負担が中心的な制約です",
        },
    },
    "Deep Learning and Representation Learning": {
        "en": {
            "name": "Deep Learning and Representation Learning",
            "focus": "CNNs, temporal models, graph networks, transformers, self-supervised learning, and cross-subject transfer",
            "challenge": "dataset leakage, weak splits, low interpretability, and poor low-data robustness can overstate performance",
        },
        "ko": {
            "name": "딥러닝 및 표현학습",
            "focus": "CNN, 시간 모델, 그래프 네트워크, 트랜스포머, 자기지도학습, 피험자 간 전이",
            "challenge": "데이터 누수, 약한 분할, 낮은 해석가능성, 저데이터 강건성 부족이 성능을 과대평가할 수 있습니다",
        },
        "zh": {
            "name": "深度学习与表征学习",
            "focus": "CNN、时序模型、图网络、Transformer、自监督学习和跨被试迁移",
            "challenge": "数据泄漏、划分不严、可解释性低和低数据鲁棒性不足可能夸大性能",
        },
        "ja": {
            "name": "深層学習と表現学習",
            "focus": "CNN、時系列モデル、グラフネットワーク、Transformer、自己教師あり学習、被験者間転移",
            "challenge": "データリーク、不十分な分割、低い解釈性、少量データでの脆弱性が性能を過大評価させる可能性があります",
        },
    },
    "EEG Signal Processing and Datasets": {
        "en": {
            "name": "EEG Signal Processing and Datasets",
            "focus": "preprocessing, artifacts, channel selection, spatial filtering, open datasets, and benchmark protocols",
            "challenge": "dataset heterogeneity and inconsistent evaluation protocols make direct comparisons fragile",
        },
        "ko": {
            "name": "EEG 신호처리 및 데이터셋",
            "focus": "전처리, 아티팩트, 채널 선택, 공간 필터링, 공개 데이터셋, 벤치마크 프로토콜",
            "challenge": "데이터셋 이질성과 평가 프로토콜 불일치가 직접 비교를 취약하게 만듭니다",
        },
        "zh": {
            "name": "EEG 信号处理与数据集",
            "focus": "预处理、伪迹、通道选择、空间滤波、开放数据集和基准协议",
            "challenge": "数据集异质性和评估协议不一致使直接比较较脆弱",
        },
        "ja": {
            "name": "EEG 信号処理とデータセット",
            "focus": "前処理、アーチファクト、チャネル選択、空間フィルタリング、公開データセット、ベンチマーク",
            "challenge": "データセットの不均一性と評価プロトコルの不一致により直接比較は脆弱になります",
        },
    },
    "Speech, Language, and Communication BCIs": {
        "en": {
            "name": "Speech, Language, and Communication BCIs",
            "focus": "spelling, text production, speech decoding, imagined speech, language models, and assistive communication",
            "challenge": "vocabulary size, latency, privacy, and autonomy remain difficult to optimize together",
        },
        "ko": {
            "name": "음성·언어·의사소통 BCI",
            "focus": "스펠링, 텍스트 생성, 음성 디코딩, 상상 음성, 언어 모델, 보조 의사소통",
            "challenge": "어휘 규모, 지연시간, 프라이버시, 사용자 자율성을 동시에 최적화하기 어렵습니다",
        },
        "zh": {
            "name": "语音、语言与通信 BCI",
            "focus": "拼写、文本生成、语音解码、想象语音、语言模型和辅助通信",
            "challenge": "词汇规模、延迟、隐私和用户自主性难以同时优化",
        },
        "ja": {
            "name": "音声・言語・コミュニケーション BCI",
            "focus": "スペリング、テキスト生成、音声デコーディング、想像音声、言語モデル、支援コミュニケーション",
            "challenge": "語彙規模、遅延、プライバシー、ユーザー自律性を同時に最適化することは困難です",
        },
    },
    "Hybrid, Affective, and Closed-loop BCIs": {
        "en": {
            "name": "Hybrid, Affective, and Closed-loop BCIs",
            "focus": "multimodal control, affective state awareness, neurofeedback, adaptive decoding, and closed-loop training",
            "challenge": "added sensors and adaptive feedback improve flexibility but complicate evaluation and usability",
        },
        "ko": {
            "name": "하이브리드·정서·폐루프 BCI",
            "focus": "다중모달 제어, 정서 상태 인식, 뉴로피드백, 적응형 디코딩, 폐루프 훈련",
            "challenge": "추가 센서와 적응형 피드백은 유연성을 높이지만 평가와 사용성을 복잡하게 만듭니다",
        },
        "zh": {
            "name": "混合、情感与闭环 BCI",
            "focus": "多模态控制、情感状态感知、神经反馈、自适应解码和闭环训练",
            "challenge": "额外传感器和自适应反馈提高灵活性，但会使评估和可用性复杂化",
        },
        "ja": {
            "name": "ハイブリッド・感情・閉ループ BCI",
            "focus": "マルチモーダル制御、感情状態認識、ニューロフィードバック、適応デコーディング、閉ループ訓練",
            "challenge": "追加センサーと適応フィードバックは柔軟性を高めますが、評価と使いやすさを複雑にします",
        },
    },
    "General BCI Methods and Systems": {
        "en": {
            "name": "General BCI Methods and Systems",
            "focus": "system architecture, conceptual frameworks, surveys, evaluation principles, ethics, usability, and translation",
            "challenge": "broad survey papers can dominate citation-ranked views while detailed empirical evidence remains scattered",
        },
        "ko": {
            "name": "일반 BCI 방법론 및 시스템",
            "focus": "시스템 아키텍처, 개념 프레임워크, 서베이, 평가 원칙, 윤리, 사용성, 전환 연구",
            "challenge": "광범위한 서베이 논문이 인용 순위를 지배하면서 세부 실증 근거가 흩어져 보일 수 있습니다",
        },
        "zh": {
            "name": "通用 BCI 方法与系统",
            "focus": "系统架构、概念框架、综述、评估原则、伦理、可用性和转化",
            "challenge": "广泛综述可能主导引用排序，而细粒度实证证据仍较分散",
        },
        "ja": {
            "name": "一般的な BCI 方法論とシステム",
            "focus": "システム構成、概念枠組み、レビュー、評価原則、倫理、ユーザビリティ、実装への移行",
            "challenge": "広範なレビュー論文が引用順位を支配し、詳細な実証知見が散在して見える可能性があります",
        },
    },
}

IMPORTANT_VENUES = [
    "nature", "science", "cell", "neuron", "nature biomedical engineering",
    "nature neuroscience", "nature communications", "science robotics",
    "journal of neural engineering", "neuroimage", "ieee transactions",
    "ieee trans", "frontiers in neuroscience", "frontiers in neurorobotics",
    "brain", "npj", "plos", "elife", "neural networks",
]


def norm_text(value):
    return re.sub(r"\s+", " ", value or "").strip()


def safe_slug(value):
    value = re.sub(r"[^a-zA-Z0-9]+", "-", value.lower()).strip("-")
    return value[:80] or "paper"


def period_chart_filename(kind, start, end):
    return f"{kind}_{start}_{end}.png"


def period_chart_src(kind, start, end):
    return f"assets/periods/{period_chart_filename(kind, start, end)}"


def taxonomy_visual_src(category):
    return f"assets/taxonomy/{safe_slug(category)}.svg"


def shields_keyword_badge(keyword, color):
    message = keyword.replace("-", "--")
    return f"![{keyword}](https://img.shields.io/badge/keyword-{message}-{color})"


def shields_keyword_badge_img(keyword):
    color = KEYWORD_COLORS.get(keyword, "64748b")
    message = keyword.replace("-", "--")
    src = f"https://img.shields.io/badge/keyword-{message}-{color}"
    return f'<img alt="{html.escape(keyword)}" src="{src}">'


def readme_keyword_badges(keywords):
    return " ".join(shields_keyword_badge_img(keyword) for keyword in keywords) or "<sub>n/a</sub>"


def readme_keyword_convention_lines():
    lines = [
        "## Keywords Convention",
        "",
        "These badges define the BCI keyword tags used to read and extend this collection.",
        "",
    ]
    for keyword, description, color in KEYWORD_CONVENTION:
        lines.append(f"- {shields_keyword_badge(keyword, color)} **{keyword}**: {description}")
    return lines


def site_keyword_convention_html():
    items = []
    for keyword, description, color in KEYWORD_CONVENTION:
        items.append(
            f"<button class='keyword-item' type='button' data-keyword='{html.escape(keyword, quote=True)}' aria-pressed='false'>"
            f"<span class='keyword-chip' style='--chip-color:#{color}'>{html.escape(keyword)}</span>"
            f"<span>{html.escape(description)}</span>"
            "</button>"
        )
    return "\n".join(items)


def taxonomy_visual_svg(category):
    visual = TAXONOMY_VISUALS.get(category, TAXONOMY_VISUALS["General BCI Methods and Systems"])
    accent = visual["accent"]
    secondary = visual["secondary"]
    motif = visual["motif"]
    title = html.escape(category)
    common_brain = f"""
      <path d="M126 81c-20 0-35 13-35 32 0 20 14 35 34 35h37c19 0 34-14 34-33 0-18-14-32-32-32-5-13-19-22-38-22z" fill="#ffffff" stroke="{accent}" stroke-width="5" stroke-linejoin="round"/>
      <path d="M117 91c-9 6-13 16-9 25m23-29c-9 8-10 20-2 30m25-26c10 8 12 20 3 31m-42 10c13 8 29 8 44 0" fill="none" stroke="{secondary}" stroke-width="4" stroke-linecap="round"/>
      <circle cx="129" cy="124" r="3" fill="{accent}"/>
      <circle cx="153" cy="124" r="3" fill="{accent}"/>
      <path d="M133 135q9 7 18 0" fill="none" stroke="{accent}" stroke-width="4" stroke-linecap="round"/>
    """
    motifs = {
        "motor": f"""
          {common_brain}
          <path d="M79 158c31-8 47-20 60-39" fill="none" stroke="{accent}" stroke-width="7" stroke-linecap="round"/>
          <path d="M74 159l-20 23m26-20 19 22m-24-25 4-34" fill="none" stroke="{secondary}" stroke-width="8" stroke-linecap="round"/>
          <path d="M198 149c25-8 44-24 55-48" fill="none" stroke="{accent}" stroke-width="5" stroke-linecap="round" stroke-dasharray="7 10"/>
          <path d="M249 96l8 3-2 9" fill="none" stroke="{accent}" stroke-width="5" stroke-linecap="round"/>
        """,
        "speller": f"""
          <rect x="67" y="55" width="185" height="123" rx="18" fill="#ffffff" stroke="{accent}" stroke-width="5"/>
          <path d="M115 188h90m-45-10v24" stroke="{accent}" stroke-width="6" stroke-linecap="round"/>
          <g fill="{secondary}">
            <rect x="91" y="78" width="32" height="25" rx="8"/>
            <rect x="143" y="78" width="32" height="25" rx="8"/>
            <rect x="195" y="78" width="32" height="25" rx="8"/>
            <rect x="91" y="119" width="32" height="25" rx="8"/>
            <rect x="143" y="119" width="32" height="25" rx="8"/>
            <rect x="195" y="119" width="32" height="25" rx="8"/>
          </g>
          <circle cx="159" cy="132" r="18" fill="#fff7ed" stroke="{accent}" stroke-width="5"/>
          <path d="M159 116l4 10 11 1-8 7 2 11-9-6-9 6 2-11-8-7 11-1z" fill="{accent}"/>
        """,
        "rehab": f"""
          <path d="M96 66v74c0 25 19 45 44 45h35c26 0 48-22 48-48v-37" fill="#ffffff" stroke="{accent}" stroke-width="6" stroke-linecap="round"/>
          <path d="M124 65v73m27-82v82m27-75v77m26-56v61" stroke="{secondary}" stroke-width="9" stroke-linecap="round"/>
          <path d="M91 132c29 2 52 16 67 42" fill="none" stroke="{accent}" stroke-width="6" stroke-linecap="round"/>
          <circle cx="230" cy="66" r="18" fill="#dcfce7" stroke="{secondary}" stroke-width="5"/>
          <path d="M221 66h18m-9-9v18" stroke="{secondary}" stroke-width="5" stroke-linecap="round"/>
        """,
        "implant": f"""
          {common_brain}
          <g stroke="{accent}" stroke-width="5" stroke-linecap="round">
            <path d="M96 61l-28-28m54 25-9-36m56 43 26-37m17 74 42-9"/>
          </g>
          <g fill="{secondary}" stroke="#ffffff" stroke-width="3">
            <circle cx="68" cy="33" r="11"/>
            <circle cx="113" cy="22" r="11"/>
            <circle cx="195" cy="28" r="11"/>
            <circle cx="254" cy="93" r="11"/>
          </g>
          <rect x="115" y="151" width="92" height="31" rx="10" fill="#eef2ff" stroke="{accent}" stroke-width="5"/>
          <path d="M130 165h62" stroke="{secondary}" stroke-width="5" stroke-linecap="round"/>
        """,
        "deep": f"""
          <g stroke="{secondary}" stroke-width="4" opacity="0.9">
            <path d="M82 74l66 38-66 48m66-48 83-39m-83 39 82 47m-82-47v69"/>
            <path d="M82 160l67 21 81-22"/>
          </g>
          <g fill="#ffffff" stroke="{accent}" stroke-width="5">
            <circle cx="82" cy="74" r="18"/>
            <circle cx="82" cy="160" r="18"/>
            <circle cx="148" cy="112" r="18"/>
            <circle cx="149" cy="181" r="18"/>
            <circle cx="231" cy="73" r="18"/>
            <circle cx="230" cy="159" r="18"/>
          </g>
          <path d="M70 74h24m-12-12v24m137 71h22m-11-11v22" stroke="{accent}" stroke-width="4" stroke-linecap="round"/>
        """,
        "eeg": f"""
          <rect x="58" y="55" width="204" height="124" rx="18" fill="#ffffff" stroke="{accent}" stroke-width="5"/>
          <path d="M77 118c18 0 12-42 31-42s11 83 32 83 13-71 34-71 13 53 34 53 12-31 34-31" fill="none" stroke="{secondary}" stroke-width="6" stroke-linecap="round"/>
          <path d="M85 189h150" stroke="{accent}" stroke-width="6" stroke-linecap="round"/>
          <path d="M108 70v95m52-95v95m52-95v95" stroke="#dbeafe" stroke-width="3"/>
          <path d="M226 58l22-28m-8 27 26-5" stroke="{accent}" stroke-width="5" stroke-linecap="round"/>
        """,
        "speech": f"""
          {common_brain}
          <path d="M203 62h49c20 0 36 15 36 34s-16 34-36 34h-20l-25 24 6-24h-10c-20 0-36-15-36-34s16-34 36-34z" fill="#ffffff" stroke="{accent}" stroke-width="5" stroke-linejoin="round"/>
          <circle cx="215" cy="96" r="5" fill="{secondary}"/>
          <circle cx="235" cy="96" r="5" fill="{secondary}"/>
          <circle cx="255" cy="96" r="5" fill="{secondary}"/>
          <path d="M79 169c24 16 58 16 82 0" fill="none" stroke="{secondary}" stroke-width="7" stroke-linecap="round"/>
        """,
        "hybrid": f"""
          {common_brain}
          <path d="M78 80c-22 33-17 79 13 108m151-6c29-34 29-80 1-113" fill="none" stroke="{accent}" stroke-width="6" stroke-linecap="round"/>
          <path d="M78 80l2 19 17-8m145 91-18-2 7-17" fill="none" stroke="{accent}" stroke-width="6" stroke-linecap="round" stroke-linejoin="round"/>
          <path d="M65 154c-21-19-5-51 20-34 25-17 41 15 20 34l-20 18z" fill="#fee2e2" stroke="{secondary}" stroke-width="5" stroke-linejoin="round"/>
          <circle cx="248" cy="62" r="22" fill="#f5f3ff" stroke="{secondary}" stroke-width="5"/>
          <path d="M248 48v14l11 7" stroke="{secondary}" stroke-width="5" stroke-linecap="round"/>
        """,
        "general": f"""
          {common_brain}
          <rect x="55" y="143" width="76" height="47" rx="10" fill="#ffffff" stroke="{accent}" stroke-width="5"/>
          <rect x="192" y="143" width="76" height="47" rx="10" fill="#ffffff" stroke="{secondary}" stroke-width="5"/>
          <path d="M132 166h58m-28-28v54" stroke="{accent}" stroke-width="5" stroke-linecap="round" stroke-dasharray="6 9"/>
          <path d="M74 160h38m-38 15h27m109-15h38m-38 15h26" stroke="#94a3b8" stroke-width="4" stroke-linecap="round"/>
        """,
    }
    motif_svg = motifs.get(motif, motifs["general"])
    return f"""<svg xmlns="http://www.w3.org/2000/svg" width="320" height="220" viewBox="0 0 320 220" role="img" aria-labelledby="title desc">
  <title id="title">{title}</title>
  <desc id="desc">Original license-safe schematic illustration for the {title} taxonomy.</desc>
  <rect width="320" height="220" rx="28" fill="#f8fafc"/>
  <circle cx="64" cy="54" r="35" fill="{secondary}" opacity="0.16"/>
  <circle cx="266" cy="176" r="41" fill="{accent}" opacity="0.12"/>
  <path d="M31 190c45-24 93-24 144 0s92 24 114 0" fill="none" stroke="#e2e8f0" stroke-width="7" stroke-linecap="round"/>
  {motif_svg}
</svg>
"""


def write_taxonomy_illustrations(categories):
    asset_dir = DOCS_DIR / "assets" / "taxonomy"
    asset_dir.mkdir(parents=True, exist_ok=True)
    for category in categories:
        target = asset_dir / f"{safe_slug(category)}.svg"
        target.write_text(taxonomy_visual_svg(category), encoding="utf-8")


def paper_key(paper):
    ext = paper.get("externalIds") or {}
    for key in ("DOI", "ArXiv", "PubMed", "CorpusId"):
        if ext.get(key):
            return f"{key}:{ext[key]}"
    return paper.get("paperId") or paper.get("url") or paper.get("title", "")


def is_relevant(paper):
    title = paper.get("title") or ""
    abstract = paper.get("abstract") or ""
    title_matches = title_relevance_count(paper)
    abstract_matches = sum(1 for pattern in RELEVANCE_PATTERNS if re.search(pattern, abstract, re.I))
    return title_matches >= 1 or abstract_matches >= 2


def title_relevance_count(paper):
    title = paper.get("title") or ""
    return sum(1 for pattern in RELEVANCE_PATTERNS if re.search(pattern, title, re.I))


def venue_name(paper):
    venue = paper.get("venue") or ""
    pv = paper.get("publicationVenue") or {}
    return norm_text(venue or pv.get("name") or "")


def author_text(paper, max_authors=6):
    authors = paper.get("authors") or []
    names = [norm_text(a.get("name", "")) for a in authors if a.get("name")]
    if len(names) > max_authors:
        return ", ".join(names[:max_authors]) + ", et al."
    return ", ".join(names)


def category_for(paper):
    text = f"{paper.get('title') or ''} {paper.get('abstract') or ''}".lower()
    for category, needles in CATEGORIES:
        if any(n in text for n in needles):
            return category
    return "General BCI Methods and Systems"


def primary_link(paper):
    ext = paper.get("externalIds") or {}
    if ext.get("DOI"):
        return f"https://doi.org/{ext['DOI']}"
    if ext.get("ArXiv"):
        return f"https://arxiv.org/abs/{ext['ArXiv']}"
    return paper.get("url") or ""


def importance_score(paper):
    """Transparent metadata score used to reduce 500 candidates to 100 papers."""
    title = paper.get("title") or ""
    abstract = paper.get("abstract") or ""
    venue = venue_name(paper)
    text = f"{title} {abstract}".lower()
    venue_l = venue.lower()
    citations = int(paper.get("citationCount") or 0)
    influential = int(paper.get("influentialCitationCount") or 0)
    score = math.log1p(citations) * 22.0 + math.log1p(influential) * 18.0
    reasons = [f"citations={citations}", f"influential={influential}"]

    if any(v in venue_l for v in IMPORTANT_VENUES):
        score += 10.0
        reasons.append("recognized venue")
    if re.search(r"\b(review|survey|meta-analysis|systematic review)\b", text):
        score += 7.0
        reasons.append("review/survey")
    if re.search(r"\b(dataset|benchmark|competition|open data)\b", text):
        score += 5.0
        reasons.append("dataset/benchmark")
    if re.search(r"\b(clinical|stroke|rehabilitation|prosthe|paralysis|patient)\b", text):
        score += 5.0
        reasons.append("clinical/rehab relevance")
    if re.search(r"\b(invasive|intracortical|implant|electrocorticography|ecog)\b", text):
        score += 5.0
        reasons.append("invasive/high-bandwidth BCI")
    if re.search(r"\b(deep learning|transformer|self-supervised|domain adaptation|foundation model)\b", text):
        score += 4.0
        reasons.append("modern ML method")
    match_count = sum(1 for pattern in RELEVANCE_PATTERNS if re.search(pattern, text, re.I))
    if match_count:
        score += min(8.0, match_count * 1.5)
        reasons.append(f"BCI term matches={match_count}")
    title_match_count = sum(1 for pattern in RELEVANCE_PATTERNS if re.search(pattern, title, re.I))
    if title_match_count:
        score += min(6.0, title_match_count * 2.0)
        reasons.append(f"title BCI matches={title_match_count}")
    if paper.get("openAccessPdf"):
        score += 1.0
        reasons.append("open PDF metadata")
    return round(score, 3), "; ".join(reasons)


def fetch_year_query(year, query, max_pages=2):
    params = {
        "query": query,
        "year": str(year),
        "fields": S2_FIELDS,
        "sort": "citationCount:desc",
    }
    out = []
    token = None
    for _ in range(max_pages):
        if token:
            params["token"] = token
        resp = requests.get(S2_BULK_URL, params=params, timeout=60)
        if resp.status_code == 429:
            time.sleep(8)
            resp = requests.get(S2_BULK_URL, params=params, timeout=60)
        resp.raise_for_status()
        data = resp.json()
        out.extend(data.get("data") or [])
        token = data.get("token")
        if not token:
            break
        time.sleep(1.0)
    return out


def collect_papers():
    selected = {}
    all_candidates = {}
    for year in YEARS:
        merged = {}
        for query in QUERIES:
            print(f"[collect] {year} :: {query}", flush=True)
            try:
                for paper in fetch_year_query(year, query):
                    if paper.get("year") != year:
                        continue
                    if not paper.get("title") or not is_relevant(paper):
                        continue
                    merged[paper_key(paper)] = paper
            except Exception as exc:
                print(f"[warn] {year} {query}: {exc}", flush=True)
            time.sleep(1.0)

        ranked = sorted(
            merged.values(),
            key=lambda p: (
                importance_score(p)[0],
                int(p.get("citationCount") or 0),
                int(p.get("influentialCitationCount") or 0),
                p.get("title") or "",
            ),
            reverse=True,
        )
        candidate_pool = ranked[:CANDIDATES_PER_YEAR]
        citation_ranked = sorted(
            candidate_pool,
            key=lambda p: (
                int(p.get("citationCount") or 0),
                int(p.get("influentialCitationCount") or 0),
                importance_score(p)[0],
                p.get("title") or "",
            ),
            reverse=True,
        )
        selected_pool = citation_ranked[:TARGET_PER_YEAR]
        all_candidates[year] = [normalize_paper(p, year, i + 1, candidate=True) for i, p in enumerate(candidate_pool)]
        selected[year] = [normalize_paper(p, year, i + 1) for i, p in enumerate(selected_pool)]
        print(
            f"[collect] {year}: selected {len(selected[year])}/{TARGET_PER_YEAR} "
            f"by citations from top {len(candidate_pool)}/{min(CANDIDATES_PER_YEAR, len(ranked))} candidates "
            f"({len(ranked)} relevant records)",
            flush=True,
        )
    return selected, all_candidates


def normalize_paper(paper, year, rank, candidate=False):
    ext = paper.get("externalIds") or {}
    oa = paper.get("openAccessPdf") or {}
    fields = paper.get("s2FieldsOfStudy") or []
    score, reasons = importance_score(paper)
    return {
        "year": year,
        "rank": rank,
        "candidateRank": rank if candidate else "",
        "selectionRank": "" if candidate else rank,
        "importanceScore": score,
        "importanceReasons": reasons,
        "paperId": paper.get("paperId", ""),
        "title": norm_text(paper.get("title", "")),
        "authors": author_text(paper),
        "venue": venue_name(paper),
        "publicationDate": paper.get("publicationDate") or "",
        "citationCount": int(paper.get("citationCount") or 0),
        "influentialCitationCount": int(paper.get("influentialCitationCount") or 0),
        "category": category_for(paper),
        "abstract": norm_text(paper.get("abstract", "")),
        "doi": ext.get("DOI", ""),
        "arxiv": ext.get("ArXiv", ""),
        "pubmed": str(ext.get("PubMed", "")),
        "url": primary_link(paper),
        "semanticScholarUrl": paper.get("url", ""),
        "openAccessPdf": oa.get("url", "") if isinstance(oa, dict) else "",
        "fieldsOfStudy": "; ".join(sorted({f.get("category", "") for f in fields if f.get("category")})),
        "candidate": candidate,
    }


def write_json_csv(selected, candidates):
    flat = [paper for year in YEARS for paper in selected.get(year, [])]
    candidate_flat = [paper for year in YEARS for paper in candidates.get(year, [])]
    DATA_DIR.mkdir(exist_ok=True)
    with (DATA_DIR / PAPERS_JSON).open("w", encoding="utf-8") as f:
        json.dump({
            "generated": date.today().isoformat(),
            "candidate_pool_per_year": CANDIDATES_PER_YEAR,
            "target_per_year": TARGET_PER_YEAR,
            "papers": flat,
        }, f, ensure_ascii=False, indent=2)
    with (DATA_DIR / CANDIDATES_JSON).open("w", encoding="utf-8") as f:
        json.dump({
            "generated": date.today().isoformat(),
            "candidate_pool_per_year": CANDIDATES_PER_YEAR,
            "candidates": candidate_flat,
        }, f, ensure_ascii=False, indent=2)
    fields = list(flat[0].keys()) if flat else []
    candidate_fields = list(candidate_flat[0].keys()) if candidate_flat else fields
    with (DATA_DIR / PAPERS_CSV).open("w", encoding="utf-8-sig", newline="") as f:
        writer = csv.DictWriter(f, fieldnames=fields)
        writer.writeheader()
        writer.writerows(flat)
    with (DATA_DIR / CANDIDATES_CSV).open("w", encoding="utf-8-sig", newline="") as f:
        writer = csv.DictWriter(f, fieldnames=candidate_fields)
        writer.writeheader()
        writer.writerows(candidate_flat)
    for year in YEARS:
        rows = selected.get(year, [])
        with (DATA_DIR / f"papers_{year}.csv").open("w", encoding="utf-8-sig", newline="") as f:
            writer = csv.DictWriter(f, fieldnames=fields)
            writer.writeheader()
            writer.writerows(rows)
        candidate_rows = candidates.get(year, [])
        with (DATA_DIR / f"candidates_top500_{year}.csv").open("w", encoding="utf-8-sig", newline="") as f:
            writer = csv.DictWriter(f, fieldnames=candidate_fields)
            writer.writeheader()
            writer.writerows(candidate_rows)
    return flat


def year_stats(flat):
    by_year = defaultdict(list)
    for p in flat:
        by_year[p["year"]].append(p)
    stats = {}
    for year, rows in by_year.items():
        stats[year] = {
            "count": len(rows),
            "citations": sum(p["citationCount"] for p in rows),
            "influential": sum(p["influentialCitationCount"] for p in rows),
            "top": rows[0] if rows else None,
        }
    return stats


def category_stats(flat):
    return Counter(p["category"] for p in flat)


def md_link(label, url):
    label = label.replace("|", "\\|")
    return f"[{label}]({url})" if url else label


def first_sentence(value, fallback="Metadata-only record; no abstract sentence was available."):
    text = norm_text(value)
    if not text:
        return fallback
    parts = re.split(r"(?<=[.!?])\s+", text)
    return parts[0][:360].rstrip()


def method_tags(p):
    text = f"{p.get('title', '')} {p.get('abstract', '')}".lower()
    tags = []
    checks = [
        ("review", r"\b(review|survey|meta-analysis|systematic review)\b"),
        ("dataset/benchmark", r"\b(dataset|benchmark|competition|open data)\b"),
        ("clinical/rehab", r"\b(clinical|stroke|rehabilitation|prosthe|patient|paralysis)\b"),
        ("invasive", r"\b(invasive|intracortical|implant|electrocorticography|ecog)\b"),
        ("deep learning", r"\b(deep learning|cnn|convolution|transformer|graph neural|self-supervised|domain adaptation)\b"),
        ("communication", r"\b(speech|language|speller|typing|communication)\b"),
        ("closed-loop", r"\b(closed-loop|closed loop|neurofeedback|adaptive)\b"),
    ]
    for label, pattern in checks:
        if re.search(pattern, text):
            tags.append(label)
    return tags[:4] or ["BCI method"]


def paper_keyword_tags(p):
    text = norm_text(" ".join([
        p.get("title", ""),
        p.get("abstract", ""),
        p.get("category", ""),
        p.get("venue", ""),
        p.get("fieldsOfStudy", ""),
    ])).lower()
    category = p.get("category", "")
    tags = []

    def add(keyword):
        if keyword not in tags:
            tags.append(keyword)

    if re.search(r"\b(invasive|intracortical|electrocorticography|ecog|implant|implanted|implantable|neuralink)\b", text) or category == "Invasive and Implantable Interfaces":
        add("invasive")
    if re.search(r"\b(non-invasive|noninvasive|eeg|electroencephalography|fnirs|nirs|meg|fmri|scalp|ssvep|p300|erp|smr)\b", text) or category in {"EEG Signal Processing and Datasets", "SSVEP, P300, and ERP Spellers"}:
        add("non-invasive")
    if re.search(r"\b(human|humans|participant|participants|patient|patients|volunteer|volunteers|subject|subjects|stroke|als|tetraplegic|paralysis)\b", text) or category in {"Rehabilitation and Neuroprosthetics", "Speech, Language, and Communication BCIs"}:
        add("human")
    if re.search(r"\b(non-human|nonhuman|animal|monkey|macaque|rat|rats|mouse|mice|primate|simulation|simulated)\b", text):
        add("non-human")
    if re.search(r"\b(smr|sensorimotor rhythm|motor imagery|mu rhythm|erd|ers|movement decoding)\b", text) or category == "Motor Imagery and Movement Decoding":
        add("SMR")
    if re.search(r"\b(ssvep|steady-state visual evoked|steady state visual evoked)\b", text):
        add("SSVEP")
    if re.search(r"\b(p300|p-300|event-related potential|event related potential|erp speller|speller)\b", text) or category == "SSVEP, P300, and ERP Spellers":
        add("P300")
    if re.search(r"\b(arm|hand|reach|reaching|cursor|trajectory|direction|directional|upper limb|upper-limb|robotic arm|movement signal)\b", text):
        add("arm-direction")

    ordered = [keyword for keyword, _, _ in KEYWORD_CONVENTION if keyword in tags]
    return ordered or ["non-invasive"]


METHOD_TAG_LOCALIZATION = {
    "en": {},
    "ko": {
        "review": "리뷰/서베이",
        "dataset/benchmark": "데이터셋/벤치마크",
        "clinical/rehab": "임상/재활",
        "invasive": "침습형",
        "deep learning": "딥러닝",
        "communication": "의사소통",
        "closed-loop": "폐루프",
        "BCI method": "BCI 방법론",
    },
    "zh": {
        "review": "综述/调查",
        "dataset/benchmark": "数据集/基准",
        "clinical/rehab": "临床/康复",
        "invasive": "侵入式",
        "deep learning": "深度学习",
        "communication": "交流通信",
        "closed-loop": "闭环",
        "BCI method": "BCI 方法",
    },
    "ja": {
        "review": "レビュー/サーベイ",
        "dataset/benchmark": "データセット/ベンチマーク",
        "clinical/rehab": "臨床/リハビリ",
        "invasive": "侵襲型",
        "deep learning": "深層学習",
        "communication": "コミュニケーション",
        "closed-loop": "閉ループ",
        "BCI method": "BCI 方法論",
    },
}

PAPER_ASSESSMENT_LOCALIZATION = {
    "ko": {
        "highCitation": "높은 인용 신호 ({value})",
        "influentialCitation": "영향력 있는 인용 신호 ({value})",
        "recognized venue": "공인된 주요 학술지/학회",
        "open-access PDF metadata": "오픈액세스 PDF 메타데이터 제공",
        "selected by citation count from the audited BCI candidate pool": "검토된 BCI 후보군에서 인용수 기준으로 선정",
    },
    "zh": {
        "highCitation": "高引用信号（{value}）",
        "influentialCitation": "高影响力引用信号（{value}）",
        "recognized venue": "公认的重要期刊/会议",
        "open-access PDF metadata": "提供开放获取 PDF 元数据",
        "selected by citation count from the audited BCI candidate pool": "从已审核的 BCI 候选池中按引用次数选出",
    },
    "ja": {
        "highCitation": "高い引用シグナル（{value}）",
        "influentialCitation": "影響力の高い引用シグナル（{value}）",
        "recognized venue": "評価の高い主要ジャーナル/会議",
        "open-access PDF metadata": "オープンアクセス PDF メタデータあり",
        "selected by citation count from the audited BCI candidate pool": "監査済み BCI 候補群から引用数に基づいて選定",
    },
}

PAPER_KEY_IDEA_TEMPLATES = {
    "ko": '"{title}"은(는) {category} 범주에서 {tags}와 관련된 BCI 연구의 핵심 흐름을 보여줍니다.',
    "zh": "《{title}》展示了 {category} 类别中与 {tags} 相关的 BCI 研究核心方向。",
    "ja": "「{title}」は、{category} 分野における {tags} に関連した BCI 研究の主要な流れを示します。",
}

RESEARCH_LIMITATIONS_BY_CATEGORY = {
    "Motor Imagery and Movement Decoding": [
        "SMR and movement decoders often need subject-specific calibration and may drift across sessions.",
        "Offline motor-imagery accuracy may not translate to reliable continuous control or rehabilitation gains.",
    ],
    "SSVEP, P300, and ERP Spellers": [
        "Visual spellers can be constrained by gaze dependence, stimulus comfort, and fatigue during long sessions.",
        "Controlled target layouts and calibration conditions may not transfer to everyday communication.",
    ],
    "Rehabilitation and Neuroprosthetics": [
        "Clinical cohorts are often small and heterogeneous, making functional effect sizes difficult to compare.",
        "Short follow-up windows can miss whether gains transfer to daily activities.",
    ],
    "Invasive and Implantable Interfaces": [
        "Surgical risk, device durability, and long-term signal stability constrain clinical translation.",
        "Small participant cohorts make headline decoding performance difficult to generalize.",
    ],
    "Deep Learning and Representation Learning": [
        "Deep models can overfit small EEG datasets without rigorous cross-subject and cross-device validation.",
        "Model interpretability and uncertainty estimates are often insufficient for clinical trust.",
    ],
    "EEG Signal Processing and Datasets": [
        "Dataset and preprocessing choices can bias comparisons across algorithms and laboratories.",
        "Artifact robustness and sensor-setup variability need stronger external validation.",
    ],
    "Speech, Language, and Communication BCIs": [
        "Communication systems must still balance accuracy, latency, vocabulary size, privacy, and user autonomy.",
        "Many decoding results rely on constrained tasks or limited participant diversity.",
    ],
    "Hybrid, Affective, and Closed-loop BCIs": [
        "Hybrid and closed-loop designs make it difficult to isolate which signal or feedback component drives the benefit.",
        "User learning, fatigue, and affect can change during closed-loop operation and confound evaluation.",
    ],
    "General BCI Methods and Systems": [
        "Broad BCI system claims often need stronger standardized evaluation across users, sessions, and devices.",
        "Laboratory performance may not transfer to everyday usability, safety, and maintenance conditions.",
    ],
}

RESEARCH_LIMITATIONS_BY_SIGNAL = {
    "invasive": "Surgical risk, device durability, and long-term signal stability constrain clinical translation.",
    "non-invasive": "Non-invasive signals are vulnerable to low signal-to-noise ratio, artifacts, and electrode setup variability.",
    "human": "Participant variability can limit generalization across age, ability, disease status, and training experience.",
    "non-human": "Non-human or simulated results require cautious translation to everyday human BCI use.",
    "SMR": "SMR and movement decoders often need subject-specific calibration and may drift across sessions.",
    "SSVEP": "Visual spellers can be constrained by gaze dependence, stimulus comfort, and fatigue during long sessions.",
    "P300": "P300 and ERP performance can depend on attention, stimulus timing, and error correction strategy.",
    "arm-direction": "Arm or cursor decoding performance may not generalize to unconstrained daily functional movements.",
}

RESEARCH_LIMITATIONS_BY_METHOD = {
    "review": "Review-level synthesis cannot resolve study quality differences, protocol bias, or reproducibility gaps.",
    "dataset/benchmark": "Benchmark value depends on standardized protocols, transparent splits, and external replication.",
    "clinical/rehab": "Clinical impact needs stronger longitudinal evidence on daily function, adherence, and adverse events.",
    "deep learning": "Deep models can overfit small EEG datasets without rigorous cross-subject and cross-device validation.",
    "communication": "Communication systems must still balance accuracy, latency, vocabulary size, privacy, and user autonomy.",
    "closed-loop": "Offline benchmark performance may not transfer to real-time closed-loop use.",
}

RESEARCH_LIMITATION_TRANSLATIONS = {
    "ko": {
        "SMR and movement decoders often need subject-specific calibration and may drift across sessions.": "SMR 및 움직임 디코더는 피험자별 보정이 필요한 경우가 많고 세션이 바뀌면 성능이 흔들릴 수 있습니다.",
        "Offline motor-imagery accuracy may not translate to reliable continuous control or rehabilitation gains.": "오프라인 운동상상 정확도가 안정적인 연속 제어나 재활 효과로 바로 이어지지 않을 수 있습니다.",
        "Visual spellers can be constrained by gaze dependence, stimulus comfort, and fatigue during long sessions.": "시각 기반 speller는 시선 의존성, 자극 편안함, 장시간 사용 피로의 제약을 받을 수 있습니다.",
        "Controlled target layouts and calibration conditions may not transfer to everyday communication.": "통제된 목표 배열과 보정 조건이 일상적인 의사소통 환경으로 그대로 전이되지 않을 수 있습니다.",
        "Clinical cohorts are often small and heterogeneous, making functional effect sizes difficult to compare.": "임상 코호트가 작고 이질적인 경우가 많아 기능적 효과 크기를 비교하기 어렵습니다.",
        "Short follow-up windows can miss whether gains transfer to daily activities.": "짧은 추적 관찰 기간은 개선 효과가 일상 활동으로 전이되는지 충분히 보여주지 못할 수 있습니다.",
        "Surgical risk, device durability, and long-term signal stability constrain clinical translation.": "수술 위험, 장치 내구성, 장기 신호 안정성이 임상 적용을 제한합니다.",
        "Small participant cohorts make headline decoding performance difficult to generalize.": "참여자 수가 적으면 보고된 높은 디코딩 성능을 일반화하기 어렵습니다.",
        "Deep models can overfit small EEG datasets without rigorous cross-subject and cross-device validation.": "엄격한 피험자 간/장치 간 검증이 없으면 딥러닝 모델이 작은 EEG 데이터셋에 과적합될 수 있습니다.",
        "Model interpretability and uncertainty estimates are often insufficient for clinical trust.": "모델 해석 가능성과 불확실성 추정이 임상적 신뢰를 얻기에 충분하지 않은 경우가 많습니다.",
        "Dataset and preprocessing choices can bias comparisons across algorithms and laboratories.": "데이터셋과 전처리 선택이 알고리즘 및 연구실 간 비교를 편향시킬 수 있습니다.",
        "Artifact robustness and sensor-setup variability need stronger external validation.": "아티팩트 견고성과 센서 설정 변화에 대한 외부 검증이 더 필요합니다.",
        "Communication systems must still balance accuracy, latency, vocabulary size, privacy, and user autonomy.": "의사소통 BCI는 정확도, 지연시간, 어휘 크기, 프라이버시, 사용자 자율성을 함께 균형 있게 다뤄야 합니다.",
        "Many decoding results rely on constrained tasks or limited participant diversity.": "많은 디코딩 결과가 제한된 과제나 제한적인 참여자 다양성에 의존합니다.",
        "Hybrid and closed-loop designs make it difficult to isolate which signal or feedback component drives the benefit.": "하이브리드 및 폐루프 설계에서는 어떤 신호나 피드백 요소가 효과를 만드는지 분리해 평가하기 어렵습니다.",
        "User learning, fatigue, and affect can change during closed-loop operation and confound evaluation.": "폐루프 사용 중 사용자 학습, 피로, 정서 상태가 변해 평가를 혼란스럽게 만들 수 있습니다.",
        "Broad BCI system claims often need stronger standardized evaluation across users, sessions, and devices.": "넓은 범위의 BCI 시스템 주장은 사용자, 세션, 장치 전반의 표준화된 평가가 더 필요합니다.",
        "Laboratory performance may not transfer to everyday usability, safety, and maintenance conditions.": "실험실 성능이 일상 사용성, 안전성, 유지관리 조건으로 그대로 전이되지 않을 수 있습니다.",
        "Non-invasive signals are vulnerable to low signal-to-noise ratio, artifacts, and electrode setup variability.": "비침습 신호는 낮은 신호대잡음비, 아티팩트, 전극 설정 변화에 취약합니다.",
        "Participant variability can limit generalization across age, ability, disease status, and training experience.": "나이, 능력, 질환 상태, 훈련 경험 차이가 일반화를 제한할 수 있습니다.",
        "Non-human or simulated results require cautious translation to everyday human BCI use.": "비인간 또는 시뮬레이션 결과는 실제 인간 BCI 사용으로 전환할 때 주의가 필요합니다.",
        "P300 and ERP performance can depend on attention, stimulus timing, and error correction strategy.": "P300 및 ERP 성능은 주의 수준, 자극 타이밍, 오류 보정 전략에 좌우될 수 있습니다.",
        "Arm or cursor decoding performance may not generalize to unconstrained daily functional movements.": "팔 또는 커서 디코딩 성능이 제약 없는 일상 기능 움직임으로 일반화되지 않을 수 있습니다.",
        "Review-level synthesis cannot resolve study quality differences, protocol bias, or reproducibility gaps.": "리뷰 수준의 종합만으로는 연구 품질 차이, 프로토콜 편향, 재현성 격차를 해결할 수 없습니다.",
        "Benchmark value depends on standardized protocols, transparent splits, and external replication.": "벤치마크의 가치는 표준화된 프로토콜, 투명한 데이터 분할, 외부 재현에 달려 있습니다.",
        "Clinical impact needs stronger longitudinal evidence on daily function, adherence, and adverse events.": "임상 효과는 일상 기능, 순응도, 이상반응에 대한 장기 근거가 더 필요합니다.",
        "Offline benchmark performance may not transfer to real-time closed-loop use.": "오프라인 벤치마크 성능이 실시간 폐루프 사용으로 그대로 전이되지 않을 수 있습니다.",
        "Recent work needs independent replication before claims are treated as stable.": "최신 연구의 주장은 안정적인 결론으로 보기 전에 독립적 재현이 필요합니다.",
    },
    "zh": {
        "SMR and movement decoders often need subject-specific calibration and may drift across sessions.": "SMR 和运动解码器通常需要针对个体校准，并可能在不同会话间发生性能漂移。",
        "Offline motor-imagery accuracy may not translate to reliable continuous control or rehabilitation gains.": "离线运动想象准确率未必能转化为可靠的连续控制或康复收益。",
        "Visual spellers can be constrained by gaze dependence, stimulus comfort, and fatigue during long sessions.": "视觉拼写器会受到注视依赖、刺激舒适度和长时间使用疲劳的限制。",
        "Controlled target layouts and calibration conditions may not transfer to everyday communication.": "受控目标布局和校准条件可能无法直接迁移到日常交流场景。",
        "Clinical cohorts are often small and heterogeneous, making functional effect sizes difficult to compare.": "临床队列常常规模较小且异质性较强，使功能效应量难以比较。",
        "Short follow-up windows can miss whether gains transfer to daily activities.": "较短的随访窗口可能无法判断收益是否迁移到日常活动。",
        "Surgical risk, device durability, and long-term signal stability constrain clinical translation.": "手术风险、设备耐久性和长期信号稳定性限制临床转化。",
        "Small participant cohorts make headline decoding performance difficult to generalize.": "参与者队列较小会使突出报道的解码性能难以泛化。",
        "Deep models can overfit small EEG datasets without rigorous cross-subject and cross-device validation.": "如果缺少严格的跨被试和跨设备验证，深度模型可能过拟合小规模 EEG 数据集。",
        "Model interpretability and uncertainty estimates are often insufficient for clinical trust.": "模型可解释性和不确定性估计往往不足以支撑临床信任。",
        "Dataset and preprocessing choices can bias comparisons across algorithms and laboratories.": "数据集和预处理选择可能使算法和实验室之间的比较产生偏差。",
        "Artifact robustness and sensor-setup variability need stronger external validation.": "对伪迹的鲁棒性和传感器设置差异仍需要更强的外部验证。",
        "Communication systems must still balance accuracy, latency, vocabulary size, privacy, and user autonomy.": "通信型 BCI 仍需平衡准确率、延迟、词汇量、隐私和用户自主性。",
        "Many decoding results rely on constrained tasks or limited participant diversity.": "许多解码结果依赖受限任务或有限的参与者多样性。",
        "Hybrid and closed-loop designs make it difficult to isolate which signal or feedback component drives the benefit.": "混合和闭环设计使得难以分离究竟是哪种信号或反馈组件带来收益。",
        "User learning, fatigue, and affect can change during closed-loop operation and confound evaluation.": "闭环运行中用户学习、疲劳和情绪会发生变化，从而干扰评估。",
        "Broad BCI system claims often need stronger standardized evaluation across users, sessions, and devices.": "广泛的 BCI 系统主张通常需要在用户、会话和设备之间进行更强的标准化评估。",
        "Laboratory performance may not transfer to everyday usability, safety, and maintenance conditions.": "实验室性能可能无法直接迁移到日常可用性、安全性和维护条件。",
        "Non-invasive signals are vulnerable to low signal-to-noise ratio, artifacts, and electrode setup variability.": "非侵入式信号容易受到低信噪比、伪迹和电极设置差异的影响。",
        "Participant variability can limit generalization across age, ability, disease status, and training experience.": "年龄、能力、疾病状态和训练经验差异会限制泛化。",
        "Non-human or simulated results require cautious translation to everyday human BCI use.": "非人类或模拟结果在迁移到日常人类 BCI 使用时需要谨慎。",
        "P300 and ERP performance can depend on attention, stimulus timing, and error correction strategy.": "P300 和 ERP 性能可能依赖注意力、刺激时序和纠错策略。",
        "Arm or cursor decoding performance may not generalize to unconstrained daily functional movements.": "手臂或光标解码性能可能无法泛化到无约束的日常功能动作。",
        "Review-level synthesis cannot resolve study quality differences, protocol bias, or reproducibility gaps.": "综述层面的综合无法解决研究质量差异、方案偏差或可复现性缺口。",
        "Benchmark value depends on standardized protocols, transparent splits, and external replication.": "基准的价值取决于标准化方案、透明的数据划分和外部复现。",
        "Clinical impact needs stronger longitudinal evidence on daily function, adherence, and adverse events.": "临床影响需要关于日常功能、依从性和不良事件的更强纵向证据。",
        "Offline benchmark performance may not transfer to real-time closed-loop use.": "离线基准性能可能无法迁移到实时闭环使用。",
        "Recent work needs independent replication before claims are treated as stable.": "近期研究的结论在被视为稳定之前需要独立复现。",
    },
    "ja": {
        "SMR and movement decoders often need subject-specific calibration and may drift across sessions.": "SMR や運動デコーダは被験者ごとの校正を必要とすることが多く、セッション間で性能が変動し得ます。",
        "Offline motor-imagery accuracy may not translate to reliable continuous control or rehabilitation gains.": "オフラインの運動イメージ精度が、安定した連続制御やリハビリ効果に直結するとは限りません。",
        "Visual spellers can be constrained by gaze dependence, stimulus comfort, and fatigue during long sessions.": "視覚型スペラーは視線依存性、刺激の快適性、長時間使用時の疲労に制約されます。",
        "Controlled target layouts and calibration conditions may not transfer to everyday communication.": "制御されたターゲット配置や校正条件は、日常的なコミュニケーションへそのまま移行しない可能性があります。",
        "Clinical cohorts are often small and heterogeneous, making functional effect sizes difficult to compare.": "臨床コホートは小規模で異質なことが多く、機能的な効果量を比較しにくいです。",
        "Short follow-up windows can miss whether gains transfer to daily activities.": "短い追跡期間では、改善が日常活動へ移行したかを見落とす可能性があります。",
        "Surgical risk, device durability, and long-term signal stability constrain clinical translation.": "手術リスク、デバイス耐久性、長期的な信号安定性が臨床応用を制約します。",
        "Small participant cohorts make headline decoding performance difficult to generalize.": "参加者数が少ないと、目立つデコード性能を一般化しにくくなります。",
        "Deep models can overfit small EEG datasets without rigorous cross-subject and cross-device validation.": "厳密な被験者間・デバイス間検証がないと、深層モデルは小規模 EEG データセットに過適合し得ます。",
        "Model interpretability and uncertainty estimates are often insufficient for clinical trust.": "モデルの解釈可能性や不確実性推定は、臨床的信頼には不十分なことが多いです。",
        "Dataset and preprocessing choices can bias comparisons across algorithms and laboratories.": "データセットや前処理の選択は、アルゴリズム間や研究室間の比較にバイアスを与え得ます。",
        "Artifact robustness and sensor-setup variability need stronger external validation.": "アーチファクトへの頑健性とセンサー設定のばらつきには、より強い外部検証が必要です。",
        "Communication systems must still balance accuracy, latency, vocabulary size, privacy, and user autonomy.": "コミュニケーション BCI は、精度、遅延、語彙規模、プライバシー、ユーザー自律性のバランスが必要です。",
        "Many decoding results rely on constrained tasks or limited participant diversity.": "多くのデコード結果は、制約された課題や限られた参加者多様性に依存しています。",
        "Hybrid and closed-loop designs make it difficult to isolate which signal or feedback component drives the benefit.": "ハイブリッドおよび閉ループ設計では、どの信号やフィードバック要素が効果を生むかを切り分けにくいです。",
        "User learning, fatigue, and affect can change during closed-loop operation and confound evaluation.": "閉ループ運用中にユーザー学習、疲労、感情状態が変化し、評価を混乱させる可能性があります。",
        "Broad BCI system claims often need stronger standardized evaluation across users, sessions, and devices.": "広範な BCI システムの主張には、ユーザー、セッション、デバイスをまたぐ標準化評価がさらに必要です。",
        "Laboratory performance may not transfer to everyday usability, safety, and maintenance conditions.": "実験室での性能は、日常の使いやすさ、安全性、保守条件へそのまま移行しない可能性があります。",
        "Non-invasive signals are vulnerable to low signal-to-noise ratio, artifacts, and electrode setup variability.": "非侵襲信号は低い信号対雑音比、アーチファクト、電極設定のばらつきに弱いです。",
        "Participant variability can limit generalization across age, ability, disease status, and training experience.": "年齢、能力、疾患状態、訓練経験の違いが一般化を制限し得ます。",
        "Non-human or simulated results require cautious translation to everyday human BCI use.": "非ヒトまたはシミュレーション結果を日常的な人間の BCI 利用へ移すには注意が必要です。",
        "P300 and ERP performance can depend on attention, stimulus timing, and error correction strategy.": "P300 や ERP の性能は、注意、刺激タイミング、エラー訂正戦略に依存し得ます。",
        "Arm or cursor decoding performance may not generalize to unconstrained daily functional movements.": "腕やカーソルのデコード性能は、制約のない日常的な機能動作へ一般化しない可能性があります。",
        "Review-level synthesis cannot resolve study quality differences, protocol bias, or reproducibility gaps.": "レビュー水準の統合だけでは、研究品質の差、プロトコルバイアス、再現性の不足は解決できません。",
        "Benchmark value depends on standardized protocols, transparent splits, and external replication.": "ベンチマークの価値は、標準化されたプロトコル、透明な分割、外部再現に依存します。",
        "Clinical impact needs stronger longitudinal evidence on daily function, adherence, and adverse events.": "臨床的影響には、日常機能、遵守率、有害事象に関するより強い縦断的証拠が必要です。",
        "Offline benchmark performance may not transfer to real-time closed-loop use.": "オフラインベンチマーク性能は、リアルタイム閉ループ利用へ移行しない可能性があります。",
        "Recent work needs independent replication before claims are treated as stable.": "最近の研究成果は、安定した主張とみなす前に独立した再現が必要です。",
    },
}


def localized_method_tags(value, language):
    translations = METHOD_TAG_LOCALIZATION.get(language, {})
    tags = [tag for tag in value.split("; ") if tag]
    return ", ".join(translations.get(tag, tag) for tag in tags) or translations.get("BCI method", "BCI method")


def assessment_items(value):
    return [item.strip() for item in value.split("; ") if item.strip()]


def localized_assessment_item(item, language):
    if language == "en":
        return item
    translations = PAPER_ASSESSMENT_LOCALIZATION.get(language, {})
    match = re.fullmatch(r"high citation signal \(([^)]+)\)", item)
    if match:
        return translations["highCitation"].format(value=match.group(1))
    match = re.fullmatch(r"influential citation signal \(([^)]+)\)", item)
    if match:
        return translations["influentialCitation"].format(value=match.group(1))
    return translations.get(item, item)


def localized_research_limitation(item, language):
    if language == "en":
        return item
    return RESEARCH_LIMITATION_TRANSLATIONS.get(language, {}).get(item, item)


def localized_paper_key_idea(p, language):
    if language == "en":
        return p["keyIdea"]
    profile = category_localization(p.get("category", "General BCI Methods and Systems"), language)
    template = PAPER_KEY_IDEA_TEMPLATES[language]
    return template.format(
        title=p.get("title") or "this paper",
        category=profile["name"],
        tags=localized_method_tags(p.get("methodTags", ""), language),
    )


def localized_paper_text(p, language):
    return {
        "keyIdea": localized_paper_key_idea(p, language),
        "strengths": "; ".join(localized_assessment_item(item, language) for item in assessment_items(p["strengths"])),
        "limitations": "; ".join(localized_research_limitation(item, language) for item in assessment_items(p["limitations"])),
    }


def localized_paper_attrs(p):
    attrs = []
    for language in LANGUAGES:
        localized = localized_paper_text(p, language)
        attrs.extend([
            f'data-key-idea-{language}="{html.escape(localized["keyIdea"], quote=True)}"',
            f'data-strengths-{language}="{html.escape(localized["strengths"], quote=True)}"',
            f'data-paper-limitations-{language}="{html.escape(localized["limitations"], quote=True)}"',
        ])
    return " ".join(attrs)


def research_limitations(p, method_tag_list, keyword_list):
    limitations = []

    def add(item):
        if item and item not in limitations:
            limitations.append(item)

    category = p.get("category", "General BCI Methods and Systems")
    for tag in method_tag_list:
        add(RESEARCH_LIMITATIONS_BY_METHOD.get(tag))
    for item in RESEARCH_LIMITATIONS_BY_CATEGORY.get(category, RESEARCH_LIMITATIONS_BY_CATEGORY["General BCI Methods and Systems"]):
        add(item)
    for keyword in keyword_list:
        add(RESEARCH_LIMITATIONS_BY_SIGNAL.get(keyword))
    if p.get("year", 0) >= END_YEAR - 1:
        add("Recent work needs independent replication before claims are treated as stable.")
    return limitations[:3]


def enrich_paper(p):
    """Add reproducible interpretation fields from title, abstract, and metadata."""
    abstract = p.get("abstract", "")
    idea = first_sentence(abstract, f"Positions {p.get('title', 'this paper')} within {p.get('category', 'BCI research')}.")
    tags = method_tags(p)
    keywords = paper_keyword_tags(p)
    strengths = []
    if p.get("citationCount", 0) >= 100:
        strengths.append(f"high citation signal ({p['citationCount']:,})")
    if p.get("influentialCitationCount", 0) >= 10:
        strengths.append(f"influential citation signal ({p['influentialCitationCount']:,})")
    if "recognized venue" in p.get("importanceReasons", ""):
        strengths.append("recognized venue")
    if p.get("openAccessPdf"):
        strengths.append("open-access PDF metadata")
    if not strengths:
        strengths.append("selected by citation count from the audited BCI candidate pool")

    enriched = dict(p)
    enriched["keyIdea"] = idea
    enriched["strengths"] = "; ".join(strengths[:3])
    enriched["limitations"] = "; ".join(research_limitations(p, tags, keywords))
    enriched["methodTags"] = "; ".join(tags)
    enriched["keywordTags"] = "; ".join(keywords)
    enriched["paperLink"] = p.get("url") or p.get("semanticScholarUrl") or ""
    return enriched


def enriched_flat(flat):
    return [enrich_paper(p) for p in flat]


def category_groups(flat):
    groups = defaultdict(list)
    for p in enriched_flat(flat):
        groups[p["category"]].append(p)
    for rows in groups.values():
        rows.sort(key=lambda p: (p["citationCount"], p["influentialCitationCount"], p["importanceScore"]), reverse=True)
    return groups


def taxonomy_trends(category):
    return TAXONOMY_TRENDS.get(category, TAXONOMY_TRENDS["General BCI Methods and Systems"])


def taxonomy_limitations(category):
    return TAXONOMY_LIMITATIONS.get(category, TAXONOMY_LIMITATIONS["General BCI Methods and Systems"])


def category_localization(category, language):
    profile = CATEGORY_LOCALIZATION.get(category, CATEGORY_LOCALIZATION["General BCI Methods and Systems"])
    return profile.get(language, profile["en"])


def period_key(start, end):
    return f"{start}-{end}"


def period_label(start, end):
    if start == START_YEAR and end == END_YEAR:
        return f"All years ({YEAR_RANGE_TEXT})"
    return str(start) if start == end else f"{start}-{end}"


def all_period_ranges():
    return [(start, end) for start in YEARS for end in range(start, END_YEAR + 1)]


def period_select_ranges():
    full_range = (START_YEAR, END_YEAR)
    return [full_range] + [range_pair for range_pair in all_period_ranges() if range_pair != full_range]


def research_overview_html():
    return """
    <section class="research-brief" id="researchBrief" aria-labelledby="research-timeline-title">
      <h2 id="research-timeline-title">Research Timeline</h2>
      <div class="timeline-copy">
        <p>2000-2026년 BCI 코퍼스는 통신·제어를 위한 초기 brain-computer interface 개념 정리에서 출발해 motor imagery, SSVEP/P300 speller, EEG 신호처리, 침습형 신경 디코딩, 재활·neuroprosthetics, 딥러닝 기반 representation learning으로 확장된 연구 지형도다. 총 2,447편의 선별 논문은 운동 의도 해석과 비침습 EEG 패러다임이 여전히 중심축이며, 그 주변으로 고대역폭 침습형 인터페이스와 임상 재활 응용이 성장해 왔음을 보여준다.</p>
        <p>인용 상위 흐름은 BCI2000, EEGNet, communication/control 리뷰처럼 공동 벤치마크와 시스템 정의를 만든 논문에 강하게 모인다. 최근 구간의 핵심 변화는 정확도 경쟁만으로는 부족하다는 점이다. 세션 간 안정성, 사용자 적응, 장기 사용성, 임상 기능 개선, 안전성과 개인정보 보호가 실제 전환의 기준으로 올라오고 있다.</p>
      </div>
      <h2>Research Insights</h2>
      <div class="research-insights">
        <article class="insight-box">
          <div class="insight-label">Motor Decoding</div>
          <h3>운동 의도 디코딩이 BCI의 중심축이다</h3>
          <p>가장 큰 taxonomy는 motor imagery와 movement decoding으로, 커서 제어, 팔·손 움직임, gait, prosthetic control이 장기간 핵심 문제로 유지된다.</p>
          <p class="insight-implication">시사점: 성능 비교는 단순 분류 정확도보다 사용자별 학습, online control, 기능적 이득을 함께 봐야 한다.</p>
        </article>
        <article class="insight-box">
          <div class="insight-label">Non-invasive EEG</div>
          <h3>비침습 BCI는 데이터셋과 전처리가 성능을 좌우한다</h3>
          <p>EEG, SSVEP, P300, ERP speller 연구는 접근성이 높지만 잡음, artifact, 세션 변화, montage 차이에 민감하다.</p>
          <p class="insight-implication">시사점: 공개 데이터셋, 전처리 표준, cross-subject 검증이 알고리즘 개선만큼 중요하다.</p>
        </article>
        <article class="insight-box">
          <div class="insight-label">Implants</div>
          <h3>침습형 인터페이스는 고성능과 확장성 사이에 있다</h3>
          <p>intracortical, ECoG, implantable interface는 높은 정보율과 정밀 제어를 제공하지만 장기 안정성, 수술 위험, 유지 비용이 병목이다.</p>
          <p class="insight-implication">시사점: 임상 가치는 decoding bandwidth와 안전한 장기 운영성을 함께 만족할 때 커진다.</p>
        </article>
        <article class="insight-box">
          <div class="insight-label">Rehabilitation</div>
          <h3>재활 BCI는 정확도에서 기능 회복으로 이동한다</h3>
          <p>stroke rehabilitation, neuroprosthetics, exoskeleton 연구는 신호 해석을 넘어 운동 회복, 피드백, 반복 훈련 효과를 평가한다.</p>
          <p class="insight-implication">시사점: 논문 평가는 offline score보다 clinical endpoint와 장기 추적 결과를 중시해야 한다.</p>
        </article>
        <article class="insight-box">
          <div class="insight-label">Deep Learning</div>
          <h3>딥러닝은 표현학습을 열었지만 일반화가 관건이다</h3>
          <p>EEGNet 이후 compact CNN, transformer, domain adaptation, self-supervised learning이 늘었지만 데이터 규모와 외부 검증은 여전히 제한적이다.</p>
          <p class="insight-implication">시사점: 새로운 모델은 subject/session transfer와 calibration cost를 명시해야 한다.</p>
        </article>
        <article class="insight-box">
          <div class="insight-label">Closed Loop</div>
          <h3>폐루프·hybrid BCI가 다음 통합 단계다</h3>
          <p>hybrid, affective, adaptive, closed-loop BCI는 신호 감지와 피드백·개입을 결합해 실제 사용 환경으로 연구 범위를 넓힌다.</p>
          <p class="insight-implication">시사점: 다음 연구 지도는 신경 신호, 사용자 상태, 환경 맥락을 함께 모델링해야 한다.</p>
        </article>
      </div>
    </section>
"""


def research_copy():
    return {
        "en": """
      <h2 id="research-timeline-title">Research Timeline</h2>
      <div class="timeline-copy">
        <p>The 2000-2026 BCI corpus traces the field from early communication and control systems to motor imagery, SSVEP/P300 spellers, EEG signal processing, invasive neural decoding, rehabilitation, neuroprosthetics, and deep representation learning. The 2,447 selected papers show that movement-intention decoding and non-invasive EEG paradigms remain the central axis, while implantable high-bandwidth systems and clinical rehabilitation applications continue to grow.</p>
        <p>Highly cited work clusters around shared systems and benchmarks such as BCI2000, EEGNet, and communication/control reviews. Recent progress is less about offline accuracy alone and more about cross-session stability, user adaptation, long-term usability, clinical functional gains, safety, and privacy.</p>
      </div>
      <h2>Research Insights</h2>
      <div class="research-insights">
        <article class="insight-box"><div class="insight-label">Motor Decoding</div><h3>Motor-intention decoding remains the core axis</h3><p>Motor imagery and movement decoding dominate the taxonomy through cursor control, arm and hand movement, gait, and prosthetic control.</p><p class="insight-implication">Implication: online control, user learning, and functional benefit should be evaluated alongside accuracy.</p></article>
        <article class="insight-box"><div class="insight-label">Non-invasive EEG</div><h3>Datasets and preprocessing shape non-invasive BCI</h3><p>EEG, SSVEP, P300, and ERP speller studies are accessible but sensitive to noise, artifacts, montage, and session drift.</p><p class="insight-implication">Implication: public datasets, preprocessing standards, and cross-subject validation are as important as model changes.</p></article>
        <article class="insight-box"><div class="insight-label">Implants</div><h3>Implantable BCI balances bandwidth and scalability</h3><p>Intracortical, ECoG, and implantable interfaces provide high information rates but face long-term stability, surgical risk, and maintenance constraints.</p><p class="insight-implication">Implication: clinical value depends on decoding bandwidth plus safe long-term operation.</p></article>
        <article class="insight-box"><div class="insight-label">Rehabilitation</div><h3>Rehabilitation BCI moves from accuracy to function</h3><p>Stroke rehabilitation, neuroprosthetics, and exoskeleton studies evaluate feedback, repeated training, and motor recovery.</p><p class="insight-implication">Implication: clinical endpoints and follow-up matter more than offline scores alone.</p></article>
        <article class="insight-box"><div class="insight-label">Deep Learning</div><h3>Deep learning opened representation learning, but generalization is the test</h3><p>Compact CNNs, transformers, domain adaptation, and self-supervised learning build on EEGNet, yet data scale and external validation remain limited.</p><p class="insight-implication">Implication: new models should report subject/session transfer and calibration cost.</p></article>
        <article class="insight-box"><div class="insight-label">Closed Loop</div><h3>Closed-loop and hybrid BCI are the next integration layer</h3><p>Hybrid, affective, adaptive, and closed-loop BCI combine sensing with feedback and intervention in more realistic use contexts.</p><p class="insight-implication">Implication: future maps need to model neural signals, user state, and environmental context together.</p></article>
      </div>
""",
        "ko": """
      <h2 id="research-timeline-title">연구 타임라인</h2>
      <div class="timeline-copy">
        <p>2000-2026년 BCI 코퍼스는 통신·제어를 위한 초기 brain-computer interface 개념 정리에서 출발해 motor imagery, SSVEP/P300 speller, EEG 신호처리, 침습형 신경 디코딩, 재활·neuroprosthetics, 딥러닝 기반 representation learning으로 확장된 연구 지형도다. 총 2,447편의 선별 논문은 운동 의도 해석과 비침습 EEG 패러다임이 여전히 중심축이며, 그 주변으로 고대역폭 침습형 인터페이스와 임상 재활 응용이 성장해 왔음을 보여준다.</p>
        <p>인용 상위 흐름은 BCI2000, EEGNet, communication/control 리뷰처럼 공동 벤치마크와 시스템 정의를 만든 논문에 강하게 모인다. 최근 구간의 핵심 변화는 정확도 경쟁만으로는 부족하다는 점이다. 세션 간 안정성, 사용자 적응, 장기 사용성, 임상 기능 개선, 안전성과 개인정보 보호가 실제 전환의 기준으로 올라오고 있다.</p>
      </div>
      <h2>연구 인사이트</h2>
      <div class="research-insights">
        <article class="insight-box"><div class="insight-label">Motor Decoding</div><h3>운동 의도 디코딩이 BCI의 중심축이다</h3><p>가장 큰 taxonomy는 motor imagery와 movement decoding으로, 커서 제어, 팔·손 움직임, gait, prosthetic control이 장기간 핵심 문제로 유지된다.</p><p class="insight-implication">시사점: 성능 비교는 단순 분류 정확도보다 사용자별 학습, online control, 기능적 이득을 함께 봐야 한다.</p></article>
        <article class="insight-box"><div class="insight-label">Non-invasive EEG</div><h3>비침습 BCI는 데이터셋과 전처리가 성능을 좌우한다</h3><p>EEG, SSVEP, P300, ERP speller 연구는 접근성이 높지만 잡음, artifact, 세션 변화, montage 차이에 민감하다.</p><p class="insight-implication">시사점: 공개 데이터셋, 전처리 표준, cross-subject 검증이 알고리즘 개선만큼 중요하다.</p></article>
        <article class="insight-box"><div class="insight-label">Implants</div><h3>침습형 인터페이스는 고성능과 확장성 사이에 있다</h3><p>intracortical, ECoG, implantable interface는 높은 정보율과 정밀 제어를 제공하지만 장기 안정성, 수술 위험, 유지 비용이 병목이다.</p><p class="insight-implication">시사점: 임상 가치는 decoding bandwidth와 안전한 장기 운영성을 함께 만족할 때 커진다.</p></article>
        <article class="insight-box"><div class="insight-label">Rehabilitation</div><h3>재활 BCI는 정확도에서 기능 회복으로 이동한다</h3><p>stroke rehabilitation, neuroprosthetics, exoskeleton 연구는 신호 해석을 넘어 운동 회복, 피드백, 반복 훈련 효과를 평가한다.</p><p class="insight-implication">시사점: 논문 평가는 offline score보다 clinical endpoint와 장기 추적 결과를 중시해야 한다.</p></article>
        <article class="insight-box"><div class="insight-label">Deep Learning</div><h3>딥러닝은 표현학습을 열었지만 일반화가 관건이다</h3><p>EEGNet 이후 compact CNN, transformer, domain adaptation, self-supervised learning이 늘었지만 데이터 규모와 외부 검증은 여전히 제한적이다.</p><p class="insight-implication">시사점: 새로운 모델은 subject/session transfer와 calibration cost를 명시해야 한다.</p></article>
        <article class="insight-box"><div class="insight-label">Closed Loop</div><h3>폐루프·hybrid BCI가 다음 통합 단계다</h3><p>hybrid, affective, adaptive, closed-loop BCI는 신호 감지와 피드백·개입을 결합해 실제 사용 환경으로 연구 범위를 넓힌다.</p><p class="insight-implication">시사점: 다음 연구 지도는 신경 신호, 사용자 상태, 환경 맥락을 함께 모델링해야 한다.</p></article>
      </div>
""",
        "zh": """
      <h2 id="research-timeline-title">研究时间线</h2>
      <div class="timeline-copy">
        <p>2000-2026 年的 BCI 语料展示了该领域从早期通信与控制系统，扩展到运动想象、SSVEP/P300 拼写器、EEG 信号处理、侵入式神经解码、康复、神经假体和深度表征学习的过程。2,447 篇入选论文表明，运动意图解码和非侵入式 EEG 范式仍是核心轴线，高带宽植入式接口和临床康复应用则持续增长。</p>
        <p>高被引论文集中在 BCI2000、EEGNet 以及 communication/control 综述等共同系统和基准上。近期进展不再只看离线准确率，也越来越重视跨会话稳定性、用户适应、长期可用性、临床功能收益、安全性和隐私。</p>
      </div>
      <h2>研究洞察</h2>
      <div class="research-insights">
        <article class="insight-box"><div class="insight-label">Motor Decoding</div><h3>运动意图解码仍是核心</h3><p>运动想象和运动解码覆盖光标控制、手臂与手部运动、步态和假体控制，是 BCI taxonomy 中最大的方向。</p><p class="insight-implication">启示：应同时评估在线控制、用户学习和功能收益，而不仅是分类准确率。</p></article>
        <article class="insight-box"><div class="insight-label">Non-invasive EEG</div><h3>数据集和预处理决定非侵入式 BCI 表现</h3><p>EEG、SSVEP、P300 和 ERP 拼写器易于部署，但对噪声、伪迹、导联配置和会话漂移非常敏感。</p><p class="insight-implication">启示：公开数据集、预处理标准和跨受试者验证与模型创新同样重要。</p></article>
        <article class="insight-box"><div class="insight-label">Implants</div><h3>植入式 BCI 在带宽和可扩展性之间取舍</h3><p>皮层内、ECoG 和植入式接口能提供高信息率，但长期稳定性、手术风险和维护成本仍是瓶颈。</p><p class="insight-implication">启示：临床价值取决于解码带宽与安全长期运行的同时满足。</p></article>
        <article class="insight-box"><div class="insight-label">Rehabilitation</div><h3>康复 BCI 从准确率走向功能恢复</h3><p>卒中康复、神经假体和外骨骼研究更关注反馈、重复训练和运动恢复。</p><p class="insight-implication">启示：临床终点和长期随访比离线分数更关键。</p></article>
        <article class="insight-box"><div class="insight-label">Deep Learning</div><h3>深度学习打开表征学习，但泛化仍是考验</h3><p>EEGNet 之后，紧凑 CNN、transformer、领域自适应和自监督学习增多，但数据规模和外部验证仍有限。</p><p class="insight-implication">启示：新模型应报告跨受试者/跨会话迁移和校准成本。</p></article>
        <article class="insight-box"><div class="insight-label">Closed Loop</div><h3>闭环与混合 BCI 是下一层整合</h3><p>混合、情感、自适应和闭环 BCI 将信号感知与反馈干预结合，贴近真实使用场景。</p><p class="insight-implication">启示：未来研究需要同时建模神经信号、用户状态和环境上下文。</p></article>
      </div>
""",
        "ja": """
      <h2 id="research-timeline-title">研究タイムライン</h2>
      <div class="timeline-copy">
        <p>2000-2026 年の BCI コーパスは、初期の通信・制御システムから、運動イメージ、SSVEP/P300 スペラー、EEG 信号処理、侵襲的神経デコーディング、リハビリテーション、神経義肢、深層表現学習へ広がる流れを示している。選定された 2,447 本の論文は、運動意図の解読と非侵襲 EEG パラダイムが中心であり続ける一方、高帯域の植込み型インターフェースと臨床リハビリ応用が伸びていることを示す。</p>
        <p>高被引用の流れは BCI2000、EEGNet、communication/control のレビューのような共有システムとベンチマークに集まる。近年はオフライン精度だけでなく、セッション間安定性、ユーザー適応、長期使用性、臨床的な機能改善、安全性、プライバシーが重要な基準になっている。</p>
      </div>
      <h2>研究インサイト</h2>
      <div class="research-insights">
        <article class="insight-box"><div class="insight-label">Motor Decoding</div><h3>運動意図デコーディングが中心軸である</h3><p>運動イメージと運動デコーディングは、カーソル制御、腕・手の動き、歩行、義肢制御を通じて最大の分類を形成する。</p><p class="insight-implication">示唆：精度だけでなく、オンライン制御、ユーザー学習、機能的利益を合わせて評価すべきである。</p></article>
        <article class="insight-box"><div class="insight-label">Non-invasive EEG</div><h3>非侵襲 BCI はデータセットと前処理に左右される</h3><p>EEG、SSVEP、P300、ERP スペラーは導入しやすいが、ノイズ、アーチファクト、モンタージュ、セッション変動に敏感である。</p><p class="insight-implication">示唆：公開データセット、前処理標準、被験者間検証はモデル改良と同じくらい重要である。</p></article>
        <article class="insight-box"><div class="insight-label">Implants</div><h3>植込み型 BCI は帯域幅と拡張性の間にある</h3><p>皮質内、ECoG、植込み型インターフェースは高い情報率をもたらすが、長期安定性、手術リスク、維持コストが課題である。</p><p class="insight-implication">示唆：臨床価値はデコード帯域と安全な長期運用の両立で決まる。</p></article>
        <article class="insight-box"><div class="insight-label">Rehabilitation</div><h3>リハビリ BCI は精度から機能回復へ移る</h3><p>脳卒中リハビリ、神経義肢、外骨格研究は、フィードバック、反復訓練、運動回復を評価する。</p><p class="insight-implication">示唆：オフラインスコアよりも臨床エンドポイントと長期追跡が重要である。</p></article>
        <article class="insight-box"><div class="insight-label">Deep Learning</div><h3>深層学習は表現学習を開いたが、汎化が試金石である</h3><p>EEGNet 以降、compact CNN、transformer、domain adaptation、self-supervised learning が増えたが、データ規模と外部検証はまだ限られる。</p><p class="insight-implication">示唆：新しいモデルは被験者間・セッション間転移とキャリブレーションコストを示す必要がある。</p></article>
        <article class="insight-box"><div class="insight-label">Closed Loop</div><h3>閉ループ・hybrid BCI が次の統合層である</h3><p>hybrid、affective、adaptive、closed-loop BCI は、信号検出とフィードバック・介入を結び、より実環境に近づく。</p><p class="insight-implication">示唆：今後は神経信号、ユーザー状態、環境文脈を同時にモデル化する必要がある。</p></article>
      </div>
""",
    }


def top_metadata_values(rows, key, limit=3):
    counts = Counter()
    for row in rows:
        value = row.get(key) or ""
        if key == "methodTags":
            for tag in value.split("; "):
                if tag:
                    counts[tag] += 1
        elif value:
            counts[value] += 1
    return [value for value, _ in counts.most_common(limit)]


def language_period_analysis(language, category, rows, start, end):
    profile = category_localization(category, language)
    count = len(rows)
    citations = sum(p["citationCount"] for p in rows)
    top = max(rows, key=lambda p: (p["citationCount"], p["influentialCitationCount"], p["title"]))
    year_counts = Counter(p["year"] for p in rows)
    active_year, active_count = year_counts.most_common(1)[0]
    tags = ", ".join(top_metadata_values(rows, "methodTags")) or profile["focus"]
    venues = ", ".join(top_metadata_values(rows, "venue")) or "mixed venues"
    top_title = top["title"]
    top_citations = top["citationCount"]

    templates = {
        "en": {
            "overview": [
                f"In {start}-{end}, {profile['name']} contains {count:,} selected papers and {citations:,} citations. The most active year is {active_year} ({active_count:,} papers), and the leading citation-ranked paper is \"{top_title}\" ({top_citations:,} citations).",
                f"The category emphasis in this period is {profile['focus']}. Frequent metadata tags include {tags}, with visible venue concentration around {venues}.",
            ],
            "limitations": [
                f"Research limitation in this period: {profile['challenge']}.",
                "Protocol-level differences in participants, tasks, sensors, online validation, and real-world transfer still need careful comparison before drawing strong cross-study conclusions.",
            ],
        },
        "ko": {
            "overview": [
                f"{start}-{end} 기간의 {profile['name']} 분류에는 선정 논문 {count:,}편과 인용 {citations:,}회가 포함됩니다. 가장 활발한 해는 {active_year}년({active_count:,}편)이며, 인용 기준 대표 논문은 \"{top_title}\"({top_citations:,}회)입니다.",
                f"이 기간의 핵심 초점은 {profile['focus']}입니다. 자주 나타나는 메타데이터 태그는 {tags}이고, 주요 venue 신호는 {venues}에 집중됩니다.",
            ],
            "limitations": [
                f"이 기간의 연구 한계: {profile['challenge']}.",
                "피험자, 과제, 센서, 온라인 검증, 실제 사용 전이 가능성의 프로토콜 차이를 비교해야 강한 결론을 내릴 수 있습니다.",
            ],
        },
        "zh": {
            "overview": [
                f"在 {start}-{end} 期间，{profile['name']} 类别包含 {count:,} 篇入选论文和 {citations:,} 次引用。最活跃年份是 {active_year} 年（{active_count:,} 篇），引用排序最高的代表论文是“{top_title}”（{top_citations:,} 次引用）。",
                f"该时期的类别重点是{profile['focus']}。常见元数据标签包括 {tags}，主要期刊/会议信号集中在 {venues}。",
            ],
            "limitations": [
                f"该时期的研究局限：{profile['challenge']}。",
                "在得出强跨研究结论之前，仍需仔细比较参与者、任务、传感器、在线验证和真实使用迁移等协议差异。",
            ],
        },
        "ja": {
            "overview": [
                f"{start}-{end} 年の {profile['name']} には、選定論文 {count:,} 本と引用 {citations:,} 件が含まれます。最も活発な年は {active_year} 年（{active_count:,} 本）で、引用順の代表論文は「{top_title}」（{top_citations:,} 件）です。",
                f"この期間の主な焦点は {profile['focus']} です。頻出するメタデータタグは {tags} で、主な掲載先シグナルは {venues} に見られます。",
            ],
            "limitations": [
                f"この期間の研究上の限界: {profile['challenge']}。",
                "参加者、課題、センサー、オンライン検証、実利用への移行に関するプロトコル差を比較してから、強い横断的結論を出す必要があります。",
            ],
        },
    }
    result = templates[language]
    result["categoryName"] = profile["name"]
    return result


def overall_period_summary(rows, start, end):
    category_counts = Counter(p["category"] for p in rows)
    category_citations = defaultdict(int)
    keyword_counts = Counter()
    year_counts = Counter()
    year_citations = defaultdict(int)
    for paper in rows:
        category_citations[paper["category"]] += paper["citationCount"]
        year_counts[paper["year"]] += 1
        year_citations[paper["year"]] += paper["citationCount"]
        for keyword in str(paper.get("keywordTags") or "").split("; "):
            if keyword:
                keyword_counts[keyword] += 1
    peak_year, peak_count = year_counts.most_common(1)[0] if year_counts else (None, 0)
    peak_citation_year = max(year_citations, key=year_citations.get) if year_citations else None
    top = max(rows, key=lambda p: (p["citationCount"], p["influentialCitationCount"], p["title"])) if rows else None
    return {
        "startYear": start,
        "endYear": end,
        "rangeLabel": str(start) if start == end else f"{start}-{end}",
        "totalPapers": len(rows),
        "activeYears": len(year_counts),
        "citationCount": sum(p["citationCount"] for p in rows),
        "categoryCount": len(category_counts),
        "topCategories": [
            {"name": category, "count": count, "citations": category_citations[category]}
            for category, count in category_counts.most_common(6)
        ],
        "topKeywords": [
            {"name": keyword, "count": count}
            for keyword, count in keyword_counts.most_common(6)
        ],
        "peakYear": peak_year,
        "peakYearCount": peak_count,
        "peakCitationYear": peak_citation_year,
        "peakCitationCount": year_citations.get(peak_citation_year, 0) if peak_citation_year else 0,
        "topPaper": {
            "title": top["title"],
            "year": top["year"],
            "category": top["category"],
            "url": top["url"],
            "citations": top["citationCount"],
        } if top else None,
    }


def overall_research_templates():
    return {
        "en": {
            "timelineTitle": "Research Timeline",
            "summary": [
                "For {range}, the BCI corpus contains {papers} selected papers across {activeYears} active years, with {citations} citations. The strongest taxonomy signals are {topCategories}, and the most active year is {peakYear} ({peakYearCount} papers).",
                "The leading citation-ranked paper is \"{topPaper}\" ({topPaperYear}, {topPaperCitations} citations) in {topPaperCategory}. Keywords such as {topKeywords} show how this period balances non-invasive EEG paradigms, invasive decoding, rehabilitation, communication, and representation learning.",
            ],
            "insightsTitle": "Research Insights",
            "insights": [
                {"label": "Motor Decoding", "title": "Motor intention remains the organizing axis", "body": "{topCategories} dominate the selected range, while the top paper highlights {topPaperCategory}.", "implication": "Implication: period comparisons should track online control, calibration cost, and functional benefit alongside offline accuracy."},
                {"label": "Signal Layer", "title": "BCI progress depends on signal quality and benchmarks", "body": "Frequent tags such as {topKeywords} reveal the paradigms and datasets that organize {range}.", "implication": "Implication: preprocessing, public datasets, and cross-session validation are as important as model changes."},
                {"label": "Clinical Translation", "title": "Rehabilitation and neuroprosthetics need functional endpoints", "body": "A citation mass of {citations} shows a broad literature, but clinical value depends on long-term usability, safety, and measurable functional gains.", "implication": "Implication: strong studies should report patient-relevant outcomes, follow-up, and deployment constraints."},
                {"label": "Implants", "title": "Implantable BCI trades bandwidth for operational risk", "body": "When invasive or implantable categories rise, the field is testing higher information rates under stability, surgery, and maintenance constraints.", "implication": "Implication: bandwidth claims should be read together with durability and care burden."},
                {"label": "Open Gaps", "title": "Citation-ranked BCI maps still need expert reading", "body": "Recent papers, small cohorts, negative results, and user diversity can be underweighted even when period-level patterns look clear.", "implication": "Implication: use this view to navigate hypotheses, then validate with full-text protocol review."},
            ],
        },
        "ko": {
            "timelineTitle": "연구 타임라인",
            "summary": [
                "{range} 기간의 BCI 코퍼스는 활성 연도 {activeYears}년에 걸쳐 선별 논문 {papers}편과 인용 {citations}회를 포함합니다. 가장 강한 taxonomy 신호는 {topCategories}이며, 논문 수가 가장 많은 해는 {peakYear}년({peakYearCount}편)입니다.",
                "인용 기준 최상위 논문은 {topPaperCategory} 분류의 \"{topPaper}\"({topPaperYear}, {topPaperCitations}회 인용)입니다. {topKeywords} 같은 키워드는 이 기간이 비침습 EEG, 침습 decoding, 재활, 의사소통, representation learning을 어떻게 균형 있게 다루는지 보여줍니다.",
            ],
            "insightsTitle": "Research Insights",
            "insights": [
                {"label": "Motor Decoding", "title": "운동 의도 해석은 여전히 중심축입니다", "body": "{topCategories}가 선택 기간을 주도하고, 최상위 논문은 {topPaperCategory} 신호를 강조합니다.", "implication": "시사점: 기간 비교에서는 offline accuracy와 함께 online control, calibration cost, functional benefit을 봐야 합니다."},
                {"label": "Signal Layer", "title": "BCI 진전은 신호 품질과 benchmark에 달려 있습니다", "body": "{topKeywords} 같은 빈도 높은 태그는 {range}를 조직하는 paradigm과 dataset을 드러냅니다.", "implication": "시사점: 전처리, 공개 dataset, cross-session 검증은 모델 변경만큼 중요합니다."},
                {"label": "Clinical Translation", "title": "재활과 neuroprosthetics는 기능적 endpoint가 필요합니다", "body": "총 {citations}회 인용은 넓은 문헌 기반을 보여주지만, 임상 가치는 장기 사용성, 안전성, 측정 가능한 기능 개선에 달려 있습니다.", "implication": "시사점: 강한 연구는 환자 관련 outcome, follow-up, deployment constraint를 함께 보고해야 합니다."},
                {"label": "Implants", "title": "침습형 BCI는 bandwidth와 운영 위험을 맞바꿉니다", "body": "침습형 또는 implantable 분류가 커질 때는 높은 정보율을 장기 안정성, 수술, 유지관리 제약 속에서 시험하는 구간입니다.", "implication": "시사점: bandwidth 주장은 durability와 care burden과 함께 읽어야 합니다."},
                {"label": "Open Gaps", "title": "인용 기반 BCI 지도에는 전문가 해석이 필요합니다", "body": "기간 패턴이 명확해 보여도 최신 논문, 소규모 cohort, 부정 결과, 사용자 다양성은 과소평가될 수 있습니다.", "implication": "시사점: 이 뷰는 가설 탐색에 쓰고, 핵심 판단은 full-text protocol review로 검증해야 합니다."},
            ],
        },
        "zh": {
            "timelineTitle": "研究时间线",
            "summary": [
                "在 {range} 期间，BCI 语料包含 {papers} 篇入选论文，覆盖 {activeYears} 个活跃年份，总引用为 {citations} 次。最强的 taxonomy 信号是 {topCategories}，论文数量峰值出现在 {peakYear} 年（{peakYearCount} 篇）。",
                "按引用排序的领先论文是 {topPaperCategory} 中的《{topPaper}》（{topPaperYear}，{topPaperCitations} 次引用）。{topKeywords} 等关键词显示该时期如何平衡非侵入式 EEG、侵入式解码、康复、沟通和表示学习。",
            ],
            "insightsTitle": "Research Insights",
            "insights": [
                {"label": "Motor Decoding", "title": "运动意图解码仍是组织轴", "body": "{topCategories} 主导所选时期，而最高引用论文突出 {topPaperCategory}。", "implication": "启示：时期比较应同时关注在线控制、校准成本、功能收益和离线准确率。"},
                {"label": "Signal Layer", "title": "BCI 进展依赖信号质量和基准", "body": "{topKeywords} 等高频标签揭示了组织 {range} 的范式和数据集。", "implication": "启示：预处理、公开数据集和跨会话验证与模型变化同样重要。"},
                {"label": "Clinical Translation", "title": "康复和神经假体需要功能终点", "body": "{citations} 次引用说明文献基础广泛，但临床价值取决于长期可用性、安全性和可测量的功能改善。", "implication": "启示：强研究应报告患者相关结局、随访和部署约束。"},
                {"label": "Implants", "title": "植入式 BCI 在带宽和运行风险之间权衡", "body": "当侵入式或植入式分类上升时，领域正在稳定性、手术和维护约束下测试更高信息率。", "implication": "启示：带宽主张应与耐久性和护理负担一起解读。"},
                {"label": "Open Gaps", "title": "引用排序的 BCI 地图仍需专家阅读", "body": "即使时期模式清晰，最新论文、小 cohort、负结果和用户多样性也可能被低估。", "implication": "启示：用此视图导航假设，再通过全文 protocol review 验证判断。"},
            ],
        },
        "ja": {
            "timelineTitle": "研究タイムライン",
            "summary": [
                "{range} の BCI コーパスには、{activeYears} の対象年にわたる選定論文 {papers} 本、引用 {citations} 件が含まれます。最も強い taxonomy 信号は {topCategories} で、論文数のピークは {peakYear} 年（{peakYearCount} 本）です。",
                "引用順で最上位の論文は {topPaperCategory} の「{topPaper}」（{topPaperYear}、{topPaperCitations} 件引用）です。{topKeywords} などのキーワードは、この期間が非侵襲 EEG、侵襲的 decoding、リハビリ、コミュニケーション、表現学習をどう扱うかを示します。",
            ],
            "insightsTitle": "Research Insights",
            "insights": [
                {"label": "Motor Decoding", "title": "運動意図 decoding は今も中心軸です", "body": "{topCategories} が選択期間を主導し、最上位論文は {topPaperCategory} を強調しています。", "implication": "示唆：期間比較では offline accuracy だけでなく、online control、calibration cost、functional benefit を見る必要があります。"},
                {"label": "Signal Layer", "title": "BCI の進展は信号品質と benchmark に依存します", "body": "{topKeywords} などの高頻度タグは、{range} を支える paradigm と dataset を示します。", "implication": "示唆：前処理、公開 dataset、cross-session 検証はモデル変更と同じくらい重要です。"},
                {"label": "Clinical Translation", "title": "リハビリと neuroprosthetics には機能 endpoint が必要です", "body": "引用 {citations} 件は広い文献基盤を示しますが、臨床価値は長期使用性、安全性、測定可能な機能改善に依存します。", "implication": "示唆：強い研究は患者関連 outcome、follow-up、deployment constraint を報告する必要があります。"},
                {"label": "Implants", "title": "植込み型 BCI は bandwidth と運用リスクを交換します", "body": "侵襲型または implantable 分類が上昇する期間は、高い情報率を安定性、手術、保守制約の中で試す時期です。", "implication": "示唆：bandwidth の主張は durability と care burden と一緒に読む必要があります。"},
                {"label": "Open Gaps", "title": "引用ベースの BCI 地図には専門的解釈が必要です", "body": "期間パターンが明確に見えても、最新論文、小規模 cohort、否定的結果、ユーザー多様性は過小評価されることがあります。", "implication": "示唆：このビューで仮説を探索し、主要判断は全文 protocol review で検証してください。"},
            ],
        },
    }


def build_period_analysis(flat):
    enriched = enriched_flat(flat)
    ranges = [
        {"key": period_key(start, end), "label": period_label(start, end), "from": start, "to": end}
        for start, end in period_select_ranges()
    ]
    analysis = {}
    for start, end in all_period_ranges():
        range_rows = [p for p in enriched if start <= p["year"] <= end]
        groups = category_groups(range_rows)
        range_entry = {}
        range_entry["__overall__"] = overall_period_summary(range_rows, start, end)
        for category, rows in groups.items():
            if not rows:
                continue
            slug = safe_slug(category)
            range_entry[slug] = {
                language: language_period_analysis(language, category, rows, start, end)
                for language in LANGUAGES
            }
        analysis[period_key(start, end)] = range_entry
    return {
        "generated": date.today().isoformat(),
        "yearRange": YEAR_RANGE_TEXT,
        "languages": LANGUAGES,
        "uiLabels": UI_LABELS,
        "ranges": ranges,
        "analysis": analysis,
    }


def write_period_analysis(flat):
    payload = build_period_analysis(flat)
    DATA_DIR.mkdir(exist_ok=True)
    DOCS_DIR.mkdir(exist_ok=True)
    (DOCS_DIR / "data").mkdir(exist_ok=True)
    for target in (DATA_DIR / PERIOD_ANALYSIS_JSON, DOCS_DIR / "data" / PERIOD_ANALYSIS_JSON):
        with target.open("w", encoding="utf-8") as f:
            json.dump(payload, f, ensure_ascii=False, separators=(",", ":"))


def write_taxonomy_dataset(flat):
    rows = []
    for category, papers in sorted(category_groups(flat).items()):
        for idx, p in enumerate(papers, 1):
            row = dict(p)
            row["taxonomyRank"] = idx
            row["categoryOverview"] = " ".join(taxonomy_trends(category))
            row["categoryLimitations"] = " ".join(taxonomy_limitations(category))
            row["researchTrend"] = row["categoryOverview"]
            rows.append(row)
    if not rows:
        return
    fields = [
        "category", "taxonomyRank", "year", "title", "authors", "venue", "publicationDate",
        "citationCount", "influentialCitationCount", "importanceScore", "categoryOverview",
        "categoryLimitations", "researchTrend", "methodTags", "keywordTags",
        "keyIdea", "strengths", "limitations", "paperLink", "semanticScholarUrl",
        "openAccessPdf", "doi", "arxiv", "pubmed", "fieldsOfStudy",
    ]
    with (DATA_DIR / TAXONOMY_CSV).open("w", encoding="utf-8-sig", newline="") as f:
        writer = csv.DictWriter(f, fieldnames=fields, extrasaction="ignore")
        writer.writeheader()
        writer.writerows(rows)


def readme_taxonomy_table(rows, total_count):
    out = [
        '<table width="100%">',
        "<colgroup>",
        '<col width="5%">',
        '<col width="21%">',
        '<col width="12%">',
        '<col width="12%">',
        '<col width="18%">',
        '<col width="8%">',
        '<col width="24%">',
        "</colgroup>",
        "<thead>",
        "<tr>",
        '<th align="right">Rank</th>',
        "<th>Paper</th>",
        "<th>Meta</th>",
        "<th>Keywords</th>",
        "<th>Key idea</th>",
        "<th>Strengths<br><sub>(high citation signal)</sub></th>",
        "<th>Limitations</th>",
        "</tr>",
        "</thead>",
        "<tbody>",
    ]
    for idx, p in enumerate(rows, 1):
        title = html.escape(p["title"])
        link = html.escape(p["paperLink"])
        paper = f'<a href="{link}">{title}</a>' if link else title
        authors = html.escape(p["authors"] or "Unknown authors")
        venue = html.escape(p["venue"] or "Unknown venue")
        keywords = readme_keyword_badges([tag for tag in p.get("keywordTags", "").split("; ") if tag])
        out.extend([
            "<tr>",
            f'<td align="right" width="5%">{idx}</td>',
            f'<td width="21%">{paper}<br><sub>{authors}</sub></td>',
            f'<td width="12%">{p["year"]}<br>{venue}<br>{p["citationCount"]:,} citations</td>',
            f'<td width="12%">{keywords}</td>',
            f'<td width="18%">{html.escape(p["keyIdea"])}</td>',
            f'<td align="right" width="8%">{p["citationCount"]:,}</td>',
            f'<td width="24%">{html.escape(p["limitations"])}</td>',
            "</tr>",
        ])
    if total_count > len(rows):
        out.extend([
            "<tr>",
            f'<td colspan="7">See the website and taxonomy CSV for all {total_count} papers.</td>',
            "</tr>",
        ])
    out.extend(["</tbody>", "</table>"])
    return out


def write_readme(flat):
    stats = year_stats(flat)
    cats = category_stats(flat)
    groups = category_groups(flat)
    lines = [
        "# Awesome BCI",
        "",
        "[![Awesome](https://awesome.re/badge-flat.svg)](https://awesome.re)",
        "",
        "A taxonomy-first, citation-ranked map of recent Brain-Computer Interface (BCI) research.",
        "",
        '<p align="center">',
        '  <a href="https://honggi82.github.io/awesome-BCI/">',
        '    <img src="https://img.shields.io/badge/Open_Interactive_Website-honggi82.github.io%2Fawesome--BCI-0f766e?style=for-the-badge" alt="Open Interactive Website">',
        "  </a>",
        "</p>",
        "",
        "> Browse the full interactive taxonomy site with period, language, keyword, chart, and paper-card filters: https://honggi82.github.io/awesome-BCI/",
        "",
        f"Generated on {date.today().isoformat()} from free public Semantic Scholar metadata. The current edition investigates up to {CANDIDATES_PER_YEAR} BCI-related candidate papers per year for {YEAR_RANGE_TEXT}, keeps an audited candidate pool, selects the top {TARGET_PER_YEAR} papers per year by citation count, and reorganizes the final {len(flat):,} papers by research taxonomy.",
        "",
        "## Project Links",
        "",
        "- Website: https://honggi82.github.io/awesome-BCI/",
        f"- Selected dataset: `data/{PAPERS_CSV}`",
        f"- Taxonomy dataset with paper-level ideas, strengths, and limitations: `data/{TAXONOMY_CSV}`",
        f"- Precomputed period and language analysis: `data/{PERIOD_ANALYSIS_JSON}`",
        f"- Candidate pool: `data/{CANDIDATES_CSV}`",
        "- English review draft: `paper/review_en.html`, `paper/review_en.docx`",
        "- Korean review draft: `paper/review_ko.html`",
        "",
        *readme_keyword_convention_lines(),
        "",
        "## Taxonomy Overview",
        "",
        f"- **Total selected papers**: {len(flat):,} papers",
    ]
    for cat, count in cats.most_common():
        lines.append(f"- **{cat}**: {count} papers")

    lines.extend(["", "## Taxonomy Collections", ""])
    for cat, _ in cats.most_common():
        rows = groups[cat]
        top = rows[:30]
        years = sorted({p["year"] for p in rows})
        citations = sum(p["citationCount"] for p in rows)
        lines.extend([
            f"### {cat}",
            "",
            f"- Papers selected: **{len(rows)}**",
            f"- Years covered: **{years[0]}-{years[-1]}**",
            f"- Citation count in selected set: **{citations:,}**",
            "- Category Overview (main research trends):",
            *[f"  - {trend}" for trend in taxonomy_trends(cat)],
            "- Limitations:",
            *[f"  - {limitation}" for limitation in taxonomy_limitations(cat)],
            "",
            "<details>",
            f"<summary><strong>Show representative papers for {cat}</strong></summary>",
            "",
        ])
        lines.extend(readme_taxonomy_table(top, len(rows)))
        lines.extend(["", "</details>", ""])

    lines.extend([
        "## Yearly Coverage",
        "",
        "| Year | Selected papers | Citation count | Top paper |",
        "| ---: | ---: | ---: | --- |",
    ])
    by_year = defaultdict(list)
    for p in flat:
        by_year[p["year"]].append(p)
    for year in YEARS:
        rows = by_year.get(year, [])
        top = rows[0] if rows else None
        top_link = md_link(top["title"], top["url"]) if top else "n/a"
        lines.append(
            f"| {year} | {len(rows)} | {stats.get(year, {}).get('citations', 0):,} | {top_link} |"
        )
    lines.extend([
        "## Method",
        "",
        "The collection uses Semantic Scholar's Academic Graph paper search. Queries combine broad BCI terms and common subfields, results are filtered to the target publication year, relevance-filtered by BCI terms in title/abstract, deduplicated by DOI/arXiv/PubMed/CorpusId/paperId, and reduced to a maximum of 500 candidates per year. Importance scoring is retained for candidate auditing and combines log-scaled citation count, log-scaled influential citation count, recognized venue signals, BCI relevance-term density, and bonuses for reviews/surveys, datasets/benchmarks, clinical or rehabilitation relevance, invasive/high-bandwidth interfaces, and modern ML methods. The final awesome list selects the top 100 papers per year by citation count from the audited candidate pool.",
        "",
        "## Caveats",
        "",
        f"- Citation counts favor older papers and may under-rank recent {END_YEAR} work.",
        "- Metadata search is not equivalent to a full systematic review of PDFs.",
        "- Some venues and publication dates are missing in upstream metadata.",
    ])
    (ROOT / "README.md").write_text("\n".join(lines) + "\n", encoding="utf-8")


def category_summary(category, rows):
    top_tags = Counter()
    for p in rows:
        for tag in p.get("methodTags", "").split("; "):
            if tag:
                top_tags[tag] += 1
    lead = rows[0] if rows else {}
    return {
        "count": len(rows),
        "years": f"{min(p['year'] for p in rows)}-{max(p['year'] for p in rows)}" if rows else "",
        "citations": sum(p["citationCount"] for p in rows),
        "top": lead.get("title", "n/a"),
        "tags": ", ".join(tag for tag, _ in top_tags.most_common(4)) or "BCI method",
        "trends": taxonomy_trends(category),
        "limitations": taxonomy_limitations(category),
    }


def paper_card(p, taxonomy_rank):
    title = html.escape(p["title"])
    link = html.escape(p["paperLink"])
    semantic = html.escape(p.get("semanticScholarUrl") or "")
    pdf = html.escape(p.get("openAccessPdf") or "")
    keywords = [tag for tag in p.get("keywordTags", "").split("; ") if tag]
    data_keywords = html.escape(" ".join(keywords), quote=True)
    keyword_badges = "".join(
        f'<span class="keyword-chip paper-keyword" style="--chip-color:#{KEYWORD_COLORS.get(keyword, "64748b")}">{html.escape(keyword)}</span>'
        for keyword in keywords
    )
    data_title = html.escape(p["title"], quote=True)
    data_venue = html.escape(p["venue"] or "Unknown venue", quote=True)
    data_limitations = html.escape(p["limitations"], quote=True)
    localized_attrs = localized_paper_attrs(p)
    paper_link = f'<a href="{link}">Paper</a>' if link else ""
    semantic_link = f'<a href="{semantic}">Semantic Scholar</a>' if semantic else ""
    pdf_link = f'<a href="{pdf}">PDF</a>' if pdf else ""
    links = " ".join(x for x in [paper_link, semantic_link, pdf_link] if x) or "No link"
    return f"""
      <article class="paper-card" data-year="{p['year']}" data-citations="{p['citationCount']}" data-keywords="{data_keywords}" data-title="{data_title}" data-venue="{data_venue}" data-limitations="{data_limitations}" {localized_attrs}>
        <div class="paper-rank">#{taxonomy_rank}</div>
        <div class="paper-body">
          <h3>{f'<a href="{link}">{title}</a>' if link else title}</h3>
          <p class="authors">{html.escape(p["authors"] or "Unknown authors")}</p>
          <div class="meta">
            <span>{p["year"]}</span>
            <span>{html.escape(p["venue"] or "Unknown venue")}</span>
            <span>{p["citationCount"]:,} citations</span>
            <span>{p["influentialCitationCount"]:,} influential</span>
            <span>score {p["importanceScore"]}</span>
          </div>
          <div class="paper-keywords" aria-label="Keywords">{keyword_badges}</div>
          <p><strong class="paper-key-idea-label">Key idea:</strong> <span class="paper-key-idea-text">{html.escape(p["keyIdea"])}</span></p>
          <div class="assessment">
            <p><strong class="paper-strengths-label">Strengths:</strong> <span class="paper-strengths-text">{html.escape(p["strengths"])}</span></p>
            <p><strong class="paper-limitations-label">Limitations:</strong> <span class="paper-limitations-text">{html.escape(p["limitations"])}</span></p>
          </div>
          <p class="tags">{html.escape(p["methodTags"])}</p>
          <p class="links">{links}</p>
        </div>
      </article>
    """


def taxonomy_section(category, rows):
    slug = safe_slug(category)
    summary = category_summary(category, rows)
    cards = "\n".join(paper_card(p, idx) for idx, p in enumerate(rows, 1))
    return f"""
    <section id="{slug}" class="taxonomy-section" data-category="{slug}">
      <details>
        <summary>
          <img class="summary-thumb" src="{taxonomy_visual_src(category)}" alt="">
          <span class="summary-title">{html.escape(category)}</span>
          <span class="category-count">{summary['count']} papers</span>
          <span class="category-years">{summary['years']}</span>
          <span class="category-citations">{summary['citations']:,} citations</span>
        </summary>
        <div class="section-intro">
          <div class="section-visual"><img src="{taxonomy_visual_src(category)}" alt="{html.escape(category)} illustration"></div>
          <p><strong>Representative emphasis:</strong> {html.escape(summary['tags'])}</p>
          <p><strong>Top-ranked paper:</strong> <span class="top-paper">{html.escape(summary['top'])}</span></p>
          <div class="insight-grid">
            <div class="insight-box">
              <strong class="overview-heading">Category Overview</strong>
              <ul class="category-overview-list">{''.join(f'<li>{html.escape(trend)}</li>' for trend in summary['trends'])}</ul>
            </div>
            <div class="insight-box limitation-box">
              <strong class="limitation-heading">Limitations</strong>
              <ul class="category-limitations-list">{''.join(f'<li>{html.escape(limitation)}</li>' for limitation in summary['limitations'])}</ul>
            </div>
          </div>
        </div>
        <div class="paper-list">{cards}</div>
      </details>
    </section>
    """


def all_taxonomy_section(rows):
    years = sorted({p["year"] for p in rows})
    citations = sum(p["citationCount"] for p in rows)
    top = max(rows, key=lambda p: p["citationCount"])
    return f"""
    <section id="all-taxonomies" class="taxonomy-section all-taxonomy-section" data-category="all-taxonomies">
      <details>
        <summary>
          <span class="all-taxonomy-thumb" aria-hidden="true">All</span>
          <span class="summary-title">All Taxonomies</span>
          <span class="category-count">{len(rows):,} papers</span>
          <span class="category-years">{years[0]}-{years[-1]}</span>
          <span class="category-citations">{citations:,} citations</span>
        </summary>
        <div class="section-intro">
          <p><strong>Representative emphasis:</strong> Complete filtered BCI taxonomy set</p>
          <p><strong>Top-ranked paper:</strong> <span class="top-paper">{html.escape(top['title'])}</span></p>
        </div>
        <div class="paper-list all-taxonomy-list"></div>
      </details>
    </section>
    """


def write_site(flat):
    DOCS_DIR.mkdir(exist_ok=True)
    (DOCS_DIR / "assets").mkdir(exist_ok=True)
    (DOCS_DIR / "data").mkdir(exist_ok=True)
    (DOCS_DIR / "paper").mkdir(exist_ok=True)
    groups = category_groups(flat)
    cats = category_stats(flat)
    write_taxonomy_illustrations([cat for cat, _ in cats.most_common()])
    total_cites = sum(p["citationCount"] for p in flat)
    start_year_options = "\n".join(
        f'<option value="{year}"{" selected" if year == min(YEARS) else ""}>{year}</option>'
        for year in YEARS
    )
    end_year_options = "\n".join(
        f'<option value="{year}"{" selected" if year == max(YEARS) else ""}>{year}</option>'
        for year in YEARS
    )
    period_options = "\n".join(
        f'<option value="{period_key(start, end)}" data-from="{start}" data-to="{end}"{" selected" if start == START_YEAR and end == END_YEAR else ""}>{period_label(start, end)}</option>'
        for start, end in period_select_ranges()
    )
    language_options = "\n".join(
        f'<option value="{code}"{" selected" if code == "en" else ""}>{html.escape(label)}</option>'
        for code, label in LANGUAGES.items()
    )
    research_copy_payload = json.dumps(research_copy(), ensure_ascii=False)
    overall_research_templates_payload = json.dumps(overall_research_templates(), ensure_ascii=False)
    year_filter_script = """
  <script>
    (() => {
      const periodSelect = document.getElementById("periodPreset");
      const languageSelect = document.getElementById("languageSelect");
      const startSelect = document.getElementById("startYear");
      const endSelect = document.getElementById("endYear");
      const resetButton = document.getElementById("resetYears");
      const rangeStatus = document.getElementById("rangeStatus");
      const statPapers = document.getElementById("statPapers");
      const statYears = document.getElementById("statYears");
      const statCitations = document.getElementById("statCitations");
      const statCategories = document.getElementById("statCategories");
      const taxonomyTotalSummary = document.getElementById("taxonomyTotalSummary");
      const keywordFilterStatus = document.getElementById("keywordFilterStatus");
      const categoryChart = document.getElementById("categoryDistributionChart");
      const citationChart = document.getElementById("yearlyCitationsChart");
      const categoryChartCaption = document.getElementById("categoryChartCaption");
      const citationChartCaption = document.getElementById("citationChartCaption");
      const defaultStart = startSelect.value;
      const defaultEnd = endSelect.value;
      const validYears = Array.from(startSelect.options).map(option => option.value);
      const periodOptions = Array.from(periodSelect.options);
      const keywordGrid = document.querySelector(".keyword-grid");
      const keywordButtons = Array.from(document.querySelectorAll(".keyword-item[data-keyword]"));
      const allTaxonomiesSection = document.querySelector(".all-taxonomy-section");
      const allTaxonomiesDetails = allTaxonomiesSection?.querySelector("details");
      const allTaxonomiesList = allTaxonomiesSection?.querySelector(".all-taxonomy-list");
      const defaultLanguage = languageSelect.value;
      const researchCopy = __RESEARCH_COPY__;
      const overallResearchTemplates = __OVERALL_RESEARCH_TEMPLATES__;
      let allTaxonomiesCards = [];
      let precomputed = null;

      function formatNumber(value) {
        return Number(value).toLocaleString("en-US");
      }

      function yearRangeText(years) {
        if (!years.length) return "No years";
        const sorted = [...new Set(years)].sort((a, b) => a - b);
        return sorted[0] === sorted[sorted.length - 1]
          ? String(sorted[0])
          : `${sorted[0]}-${sorted[sorted.length - 1]}`;
      }

      function selectedRangeValue(start, end) {
        const match = periodOptions.find(option =>
          option.dataset.from === String(start) && option.dataset.to === String(end)
        );
        return match ? match.value : "custom";
      }

      function updatePeriodSelect(start, end) {
        periodSelect.value = selectedRangeValue(start, end);
      }

      function setFromUrl() {
        const params = new URLSearchParams(window.location.search);
        const lang = params.get("lang");
        if (lang && Array.from(languageSelect.options).some(option => option.value === lang)) {
          languageSelect.value = lang;
        }
        const requestedKeywords = (params.get("keywords") || "").split(",").filter(Boolean);
        const keywordValues = keywordButtons.map(button => button.dataset.keyword);
        const activeKeyword = requestedKeywords.find(keyword => keywordValues.includes(keyword)) || "";
        keywordButtons.forEach(button => {
          setKeywordPressed(button, button.dataset.keyword === activeKeyword);
        });
        const period = params.get("period");
        if (period) {
          const option = periodOptions.find(item => item.value === period && item.dataset.from && item.dataset.to);
          if (option) {
            periodSelect.value = period;
            startSelect.value = option.dataset.from;
            endSelect.value = option.dataset.to;
            return;
          }
        }
        const from = params.get("from");
        const to = params.get("to");
        if (validYears.includes(from)) startSelect.value = from;
        if (validYears.includes(to)) endSelect.value = to;
        updatePeriodSelect(startSelect.value, endSelect.value);
      }

      function syncUrl(start, end) {
        const url = new URL(window.location.href);
        const language = languageSelect.value;
        if (language === defaultLanguage) {
          url.searchParams.delete("lang");
        } else {
          url.searchParams.set("lang", language);
        }
        if (String(start) === defaultStart && String(end) === defaultEnd) {
          url.searchParams.delete("period");
          url.searchParams.delete("from");
          url.searchParams.delete("to");
        } else {
          const periodValue = selectedRangeValue(start, end);
          if (periodValue !== "custom") {
            url.searchParams.set("period", periodValue);
            url.searchParams.delete("from");
            url.searchParams.delete("to");
          } else {
            url.searchParams.delete("period");
            url.searchParams.set("from", start);
            url.searchParams.set("to", end);
          }
        }
        const keywords = selectedKeywords();
        if (keywords.length) {
          url.searchParams.set("keywords", keywords.join(","));
        } else {
          url.searchParams.delete("keywords");
        }
        window.history.replaceState(null, "", url);
      }

      function rangeKey(start, end) {
        return `${start}-${end}`;
      }

      function rangeLabel(start, end) {
        return start === end ? String(start) : `${start}-${end}`;
      }

      function chartPath(kind, start, end) {
        return `assets/periods/${kind}_${start}_${end}.png`;
      }

      function updateCharts(start, end) {
        const label = rangeLabel(start, end);
        if (categoryChart) {
          categoryChart.src = chartPath("category_distribution", start, end);
          categoryChart.alt = `Category distribution chart for ${label}`;
        }
        if (citationChart) {
          citationChart.src = chartPath("yearly_citations", start, end);
          citationChart.alt = `Yearly citation chart for ${label}`;
        }
        if (categoryChartCaption) categoryChartCaption.textContent = `Category distribution (${label})`;
        if (citationChartCaption) citationChartCaption.textContent = `Yearly citation mass (${label})`;
      }

      function labels() {
        const fallback = {
          papers: "papers",
          categories: "categories",
          overview: "Category Overview",
          limitations: "Limitations",
          analysis: "Selected-period analysis",
          totalSelected: "Total selected papers",
          categoryCount: "Categories",
          keyIdea: "Key idea",
          strengths: "Strengths",
          paperLimitations: "Limitations"
        };
        return {
          ...fallback,
          ...(precomputed?.uiLabels?.en || {}),
          ...(precomputed?.uiLabels?.[languageSelect.value] || {})
        };
      }
      function escapeHtml(value) {
        const escapeMap = { "&": "&amp;", "<": "&lt;", ">": "&gt;", '"': "&quot;", "'": "&#39;" };
        return String(value ?? "").replace(/[&<>"']/g, ch => escapeMap[ch]);
      }
      function names(items, key = "name") {
        return (items || []).slice(0, 3).map(item => item[key]).filter(Boolean).join(", ") || "n/a";
      }
      function researchTemplateData(metric) {
        const topCategories = metric.topCategories || [];
        const topKeywords = metric.topKeywords || [];
        const topCategory = topCategories[0] || {};
        const topPaper = metric.topPaper || {};
        return {
          range: metric.rangeLabel || `${metric.startYear}-${metric.endYear}`,
          papers: formatNumber(metric.totalPapers),
          activeYears: formatNumber(metric.activeYears),
          citations: formatNumber(metric.citationCount),
          topCategories: names(topCategories),
          topCategory: topCategory.name || "n/a",
          topCategoryCount: formatNumber(topCategory.count || 0),
          topKeywords: names(topKeywords),
          peakYear: metric.peakYear || "n/a",
          peakYearCount: formatNumber(metric.peakYearCount || 0),
          peakCitationYear: metric.peakCitationYear || "n/a",
          topPaper: topPaper.title || "n/a",
          topPaperYear: topPaper.year || "n/a",
          topPaperCategory: topPaper.category || "n/a",
          topPaperCitations: formatNumber(topPaper.citations || 0)
        };
      }
      function applyTemplate(template, data) {
        let output = template || "";
        Object.keys(data).forEach(key => {
          output = output.split("{" + key + "}").join(escapeHtml(data[key]));
        });
        return output;
      }
      function renderOverallResearch(metric) {
        const copy = overallResearchTemplates[languageSelect.value] || overallResearchTemplates.en;
        const data = researchTemplateData(metric);
        const summaryHtml = (copy.summary || []).map(text => `<p>${applyTemplate(text, data)}</p>`).join("");
        const insightHtml = (copy.insights || []).map(item => `
          <article class="insight-box">
            <div class="insight-label">${escapeHtml(item.label)}</div>
            <h3>${applyTemplate(item.title, data)}</h3>
            <p>${applyTemplate(item.body, data)}</p>
            <p class="insight-implication">${applyTemplate(item.implication, data)}</p>
          </article>`).join("");
        return `
          <h2 id="research-timeline-title">${escapeHtml(copy.timelineTitle)}</h2>
          <div class="timeline-copy">${summaryHtml}</div>
          <h2>${escapeHtml(copy.insightsTitle)}</h2>
          <div class="research-insights">${insightHtml}</div>`;
      }
      function updateResearchCopy(start, end) {
        const brief = document.getElementById("researchBrief");
        if (!brief) return;
        const metric = precomputed?.analysis?.[rangeKey(start, end)]?.__overall__;
        brief.innerHTML = metric ? renderOverallResearch(metric) : (researchCopy[languageSelect.value] || researchCopy.en);
      }

      function setList(target, items) {
        if (!target || !items) return;
        target.innerHTML = "";
        items.forEach(item => {
          const li = document.createElement("li");
          li.textContent = item;
          target.appendChild(li);
        });
      }

      function selectedKeywords() {
        return keywordButtons
          .filter(button => button.getAttribute("aria-pressed") === "true")
          .map(button => button.dataset.keyword);
      }

      function setKeywordPressed(button, pressed) {
        button.setAttribute("aria-pressed", pressed ? "true" : "false");
        button.classList.toggle("is-selected", pressed);
      }

      function keywordMatches(card, selected) {
        if (!selected.length) return true;
        const cardKeywords = (card.dataset.keywords || "").split(" ").filter(Boolean);
        return selected.some(keyword => cardKeywords.includes(keyword));
      }

      function updateKeywordFilterStatus(selected, totalPapers, copy) {
        if (!keywordFilterStatus) return;
        if (selected.length) {
          keywordFilterStatus.textContent = `Selected keyword: ${selected[0]} | Matching papers: ${formatNumber(totalPapers)} ${copy.papers}`;
        } else {
          keywordFilterStatus.textContent = `Selected keyword: all | Matching papers: ${formatNumber(totalPapers)} ${copy.papers}`;
        }
      }

      function renderAllTaxonomiesCards() {
        if (!allTaxonomiesList || !allTaxonomiesDetails?.open) return;
        allTaxonomiesList.innerHTML = "";
        const fragment = document.createDocumentFragment();
        const sortedCards = [...allTaxonomiesCards].sort((a, b) => Number(b.dataset.citations || 0) - Number(a.dataset.citations || 0));
        sortedCards.forEach(card => fragment.appendChild(card.cloneNode(true)));
        allTaxonomiesList.appendChild(fragment);
      }

      function updateAllTaxonomiesSection(cards, totalPapers, years, totalCitations, copy) {
        if (!allTaxonomiesSection) return;
        allTaxonomiesCards = cards;
        allTaxonomiesSection.hidden = totalPapers === 0;
        allTaxonomiesSection.querySelector(".category-count").textContent = `${formatNumber(totalPapers)} ${copy.papers}`;
        allTaxonomiesSection.querySelector(".category-years").textContent = yearRangeText(years);
        allTaxonomiesSection.querySelector(".category-citations").textContent = `${formatNumber(totalCitations)} citations`;
        const topPaperTarget = allTaxonomiesSection.querySelector(".top-paper");
        const topPaper = [...cards].sort((a, b) => Number(b.dataset.citations || 0) - Number(a.dataset.citations || 0))[0]?.querySelector("h3");
        if (topPaper && topPaperTarget) topPaperTarget.textContent = topPaper.textContent.trim();
        if (allTaxonomiesDetails?.open) renderAllTaxonomiesCards();
        else if (allTaxonomiesList) allTaxonomiesList.innerHTML = "";
      }

      function localizedCardText(card, field, language) {
        const suffix = language.charAt(0).toUpperCase() + language.slice(1);
        return card.dataset[`${field}${suffix}`] || card.dataset[`${field}En`] || "";
      }

      function applyPaperLocalization(card, copy) {
        const language = languageSelect.value;
        [
          ["keyIdea", ".paper-key-idea-label", ".paper-key-idea-text", copy.keyIdea],
          ["strengths", ".paper-strengths-label", ".paper-strengths-text", copy.strengths],
          ["paperLimitations", ".paper-limitations-label", ".paper-limitations-text", copy.paperLimitations]
        ].forEach(([field, labelSelector, textSelector, label]) => {
          const labelNode = card.querySelector(labelSelector);
          const textNode = card.querySelector(textSelector);
          if (labelNode) labelNode.textContent = `${label}:`;
          if (textNode) textNode.textContent = localizedCardText(card, field, language);
        });
      }

      function applyPrecomputedAnalysis(section, start, end) {
        const language = languageSelect.value;
        const entry = precomputed?.analysis?.[rangeKey(start, end)]?.[section.dataset.category];
        const analysis = entry?.[language] || entry?.en;
        if (!analysis) return;
        const copy = labels();
        const title = section.querySelector(".summary-title");
        if (title) title.textContent = analysis.categoryName;
        const overviewHeading = section.querySelector(".overview-heading");
        const limitationHeading = section.querySelector(".limitation-heading");
        if (overviewHeading) overviewHeading.textContent = copy.overview;
        if (limitationHeading) limitationHeading.textContent = copy.limitations;
        setList(section.querySelector(".category-overview-list"), analysis.overview);
        setList(section.querySelector(".category-limitations-list"), analysis.limitations);
      }

      function applyYearFilter(sync = true) {
        let start = Number(startSelect.value);
        let end = Number(endSelect.value);
        const copy = labels();
        if (start > end) {
          const previousStart = start;
          start = end;
          end = previousStart;
          startSelect.value = String(start);
          endSelect.value = String(end);
        }
        updateResearchCopy(start, end);

        let totalPapers = 0;
        let totalCitations = 0;
        let activeCategories = 0;
        const activeYears = [];
        const visibleCards = [];
        const activeKeywords = selectedKeywords();

        document.querySelectorAll(".taxonomy-section:not(.all-taxonomy-section)").forEach(section => {
          let sectionCount = 0;
          let sectionCitations = 0;
          const sectionYears = [];

          section.querySelectorAll(".paper-card").forEach(card => {
            const year = Number(card.dataset.year);
            const citations = Number(card.dataset.citations || 0);
            const visible = year >= start && year <= end && keywordMatches(card, activeKeywords);
            applyPaperLocalization(card, copy);
            card.hidden = !visible;
            if (visible) {
              sectionCount += 1;
              sectionCitations += citations;
              sectionYears.push(year);
              activeYears.push(year);
              visibleCards.push(card);
            }
          });

          const hasPapers = sectionCount > 0;
          section.hidden = !hasPapers;
          if (!hasPapers) return;

          activeCategories += 1;
          totalPapers += sectionCount;
          totalCitations += sectionCitations;
          section.querySelector(".category-count").textContent = `${formatNumber(sectionCount)} ${copy.papers}`;
          section.querySelector(".category-years").textContent = yearRangeText(sectionYears);
          section.querySelector(".category-citations").textContent = `${formatNumber(sectionCitations)} citations`;
          const topPaper = section.querySelector(".paper-card:not([hidden]) h3");
          const topPaperTarget = section.querySelector(".top-paper");
          if (topPaper && topPaperTarget) topPaperTarget.textContent = topPaper.textContent.trim();
          applyPrecomputedAnalysis(section, start, end);
        });
        updateAllTaxonomiesSection(visibleCards, totalPapers, activeYears, totalCitations, copy);

        statPapers.textContent = formatNumber(totalPapers);
        statYears.textContent = formatNumber(new Set(activeYears).size);
        statCitations.textContent = formatNumber(totalCitations);
        statCategories.textContent = formatNumber(activeCategories);
        updatePeriodSelect(start, end);
        updateCharts(start, end);
        if (taxonomyTotalSummary) {
          taxonomyTotalSummary.innerHTML = `<strong>${copy.totalSelected}:</strong> ${formatNumber(totalPapers)} ${copy.papers}; <strong>${copy.categoryCount}:</strong> ${formatNumber(activeCategories)} ${copy.categories}.`;
        }
        updateKeywordFilterStatus(activeKeywords, totalPapers, copy);
        const keywordText = activeKeywords.length ? ` · ${activeKeywords.join(", ")}` : "";
        rangeStatus.textContent = `${start}-${end} · ${formatNumber(totalPapers)} ${copy.papers} · ${formatNumber(activeCategories)} ${copy.categories}${keywordText}`;
        if (sync) syncUrl(start, end);
      }

      setFromUrl();
      applyYearFilter(false);
      fetch("data/__ANALYSIS_JSON__")
        .then(response => response.json())
        .then(data => {
          precomputed = data;
          applyYearFilter(false);
        })
        .catch(() => {
          rangeStatus.textContent = "Precomputed analysis could not be loaded.";
        });
      periodSelect.addEventListener("change", () => {
        const option = periodSelect.selectedOptions[0];
        if (option && option.dataset.from && option.dataset.to) {
          startSelect.value = option.dataset.from;
          endSelect.value = option.dataset.to;
        }
        applyYearFilter(true);
      });
      languageSelect.addEventListener("change", () => applyYearFilter(true));
      startSelect.addEventListener("change", () => applyYearFilter(true));
      endSelect.addEventListener("change", () => applyYearFilter(true));
      if (keywordGrid) {
        keywordGrid.addEventListener("click", event => {
          const button = event.target.closest(".keyword-item[data-keyword]");
          if (!button || !keywordGrid.contains(button)) return;
          event.preventDefault();
          const pressed = button.getAttribute("aria-pressed") === "true";
          keywordButtons.forEach(item => setKeywordPressed(item, !pressed && item === button));
          applyYearFilter(true);
        });
      }
      allTaxonomiesDetails?.addEventListener("toggle", () => {
        if (allTaxonomiesDetails.open) renderAllTaxonomiesCards();
        else if (allTaxonomiesList) allTaxonomiesList.innerHTML = "";
      });
      resetButton.addEventListener("click", () => {
        startSelect.value = defaultStart;
        endSelect.value = defaultEnd;
        periodSelect.value = `${defaultStart}-${defaultEnd}`;
        keywordButtons.forEach(button => setKeywordPressed(button, false));
        applyYearFilter(true);
      });
    })();
  </script>
""".replace("__ANALYSIS_JSON__", PERIOD_ANALYSIS_JSON).replace("__RESEARCH_COPY__", research_copy_payload).replace("__OVERALL_RESEARCH_TEMPLATES__", overall_research_templates_payload)
    sections = [all_taxonomy_section(flat).strip()]
    for cat, _ in cats.most_common():
        sections.append(taxonomy_section(cat, groups[cat]).strip())
    keyword_convention = site_keyword_convention_html()
    research_overview = research_overview_html().strip()
    html_doc = f"""<!doctype html>
<html lang="en">
<head>
  <meta charset="utf-8">
  <meta name="viewport" content="width=device-width, initial-scale=1">
  <link rel="icon" href="data:,">
  <title>Awesome BCI</title>
  <style>
    :root {{ color-scheme: light; --ink:#18212f; --muted:#5b6678; --line:#d9dee8; --accent:#0f766e; --accent2:#7c3aed; --bg:#f7f9fc; --panel:#ffffff; }}
    body {{ margin:0; font-family: Inter, Segoe UI, Arial, sans-serif; color:var(--ink); background:var(--bg); }}
    header {{ padding:54px 7vw 34px; background:linear-gradient(120deg,#e9fbf8,#eef2ff); border-bottom:1px solid var(--line); }}
    h1 {{ font-size:48px; margin:0 0 12px; letter-spacing:0; }}
    h2 {{ margin-top:36px; }}
    p {{ line-height:1.65; color:var(--muted); }}
    main {{ padding:28px 7vw 72px; }}
    .stats {{ display:grid; grid-template-columns:repeat(auto-fit,minmax(190px,1fr)); gap:12px; margin:24px 0; }}
    .filters {{ display:flex; flex-wrap:wrap; align-items:end; gap:12px; margin:24px 0; padding:14px; background:white; border:1px solid var(--line); border-radius:8px; }}
    .filter-field {{ display:grid; gap:6px; }}
    .wide-field {{ min-width:min(100%, 280px); }}
    .filter-field label {{ font-weight:700; font-size:13px; color:#344255; }}
    select {{ min-width:104px; height:38px; border:1px solid var(--line); border-radius:8px; background:white; color:var(--ink); padding:0 10px; font:inherit; }}
    #periodPreset {{ min-width:280px; }}
    button {{ height:38px; border:1px solid #bfc8d8; border-radius:8px; background:#f6f8fb; color:var(--ink); padding:0 14px; font-weight:700; cursor:pointer; }}
    button:hover {{ background:#eef2f7; }}
    #rangeStatus {{ color:var(--muted); font-weight:700; min-height:38px; display:inline-flex; align-items:center; }}
    .figures {{ display:grid; grid-template-columns:repeat(auto-fit,minmax(280px,1fr)); gap:16px; margin:24px 0; }}
    .chart-figure {{ margin:0; }}
    .chart-figure figcaption {{ margin-top:8px; color:var(--muted); font-size:13px; font-weight:700; }}
    .figures img {{ width:100%; aspect-ratio:16 / 9; object-fit:contain; background:white; border:1px solid var(--line); border-radius:8px; display:block; }}
    .stat, .card {{ background:white; border:1px solid var(--line); border-radius:8px; padding:16px; }}
    .stat strong {{ display:block; font-size:28px; color:var(--accent); }}
    .card span {{ display:block; margin-top:8px; color:var(--muted); }}
    .research-brief {{ margin:28px 0; padding:24px 0; border-top:1px solid var(--line); border-bottom:1px solid var(--line); }}
    .timeline-copy {{ max-width:1080px; }}
    .research-insights {{ display:grid; grid-template-columns:repeat(auto-fit,minmax(260px,1fr)); gap:12px; margin-top:12px; }}
    .research-insights .insight-label {{ color:var(--accent); font-size:12px; font-weight:800; text-transform:uppercase; letter-spacing:0; }}
    .research-insights .insight-box h3 {{ margin:6px 0 8px; font-size:17px; }}
    .research-insights .insight-box p {{ margin:8px 0 0; }}
    .insight-implication {{ color:var(--ink); font-weight:700; }}
    .keyword-section {{ margin:28px 0; }}
    .keyword-grid {{ display:grid; grid-template-columns:repeat(auto-fit,minmax(240px,1fr)); gap:10px; }}
    .keyword-item {{ display:flex; gap:10px; align-items:flex-start; height:auto; min-height:54px; padding:12px; background:white; border:1px solid var(--line); border-radius:8px; color:var(--muted); line-height:1.45; text-align:left; font:inherit; cursor:pointer; }}
    .keyword-item:hover {{ background:#f8fafc; }}
    .keyword-item[aria-pressed="true"], .keyword-item.is-selected {{ border-color:var(--accent); box-shadow:0 0 0 2px rgba(15,118,110,0.16); color:var(--ink); }}
    .keyword-chip {{ flex:0 0 auto; min-width:96px; text-align:center; background:var(--chip-color); color:white; border-radius:999px; padding:4px 9px; font-size:13px; font-weight:800; }}
    .keyword-filter-status {{ margin:10px 0 0; font-weight:700; color:var(--accent); }}
    .overview-count {{ font-weight:800; color:var(--accent); }}
    nav a {{ display:inline-block; margin:0 12px 10px 0; color:var(--accent2); font-weight:600; }}
    .card {{ display:block; color:var(--ink); }}
    .taxonomy-section {{ margin-top:16px; }}
    .taxonomy-section[hidden], .paper-card[hidden] {{ display:none !important; }}
    details {{ background:var(--panel); border:1px solid var(--line); border-radius:8px; overflow:hidden; }}
    summary {{ cursor:pointer; display:grid; grid-template-columns:64px minmax(260px,1fr) repeat(3, minmax(110px, auto)); gap:12px; align-items:center; padding:14px 18px; font-weight:700; }}
    .summary-thumb, .all-taxonomy-thumb {{ width:56px; height:40px; object-fit:cover; border:1px solid var(--line); border-radius:6px; background:#f8fafc; }}
    .all-taxonomy-thumb {{ display:inline-flex; align-items:center; justify-content:center; color:var(--accent); font-weight:800; }}
    .summary-title {{ color:var(--accent); }}
    .section-intro {{ padding:0 18px 14px; border-top:1px solid var(--line); }}
    .section-visual {{ margin:14px 0 4px; }}
    .section-visual img {{ width:min(320px, 100%); aspect-ratio:16 / 11; object-fit:contain; border:1px solid var(--line); border-radius:8px; background:#f8fafc; display:block; }}
    .insight-grid {{ display:grid; grid-template-columns:repeat(auto-fit,minmax(260px,1fr)); gap:12px; margin-top:12px; }}
    .insight-box {{ padding:12px 14px; background:#f4faf8; border:1px solid #cfe7df; border-radius:8px; }}
    .limitation-box {{ background:#fff8f1; border-color:#ead7c1; }}
    .insight-box strong span {{ color:var(--muted); font-weight:600; }}
    .insight-box ul {{ margin:8px 0 0; padding-left:20px; color:var(--muted); line-height:1.55; }}
    .paper-list {{ display:grid; gap:12px; padding:16px; background:#f9fbfd; }}
    .paper-card {{ display:grid; grid-template-columns:56px 1fr; gap:14px; padding:16px; background:white; border:1px solid var(--line); border-radius:8px; }}
    .paper-rank {{ font-weight:800; color:var(--accent2); }}
    .paper-card h3 {{ margin:0 0 6px; font-size:18px; line-height:1.35; }}
    .authors {{ margin:0 0 8px; }}
    .meta {{ display:flex; flex-wrap:wrap; gap:8px; margin:8px 0 10px; }}
    .meta span, .tags {{ display:inline-block; background:#eef2f7; border:1px solid #dce3ee; border-radius:999px; padding:5px 9px; color:#344255; font-size:13px; }}
    .paper-keywords {{ display:flex; flex-wrap:wrap; gap:6px; margin:0 0 10px; }}
    .paper-keyword {{ min-width:0; font-size:12px; padding:3px 8px; }}
    .assessment {{ display:grid; grid-template-columns:repeat(auto-fit,minmax(260px,1fr)); gap:8px; }}
    .links a {{ margin-right:12px; font-weight:700; }}
    a {{ color:#0f5f97; text-decoration:none; }}
    a:hover {{ text-decoration:underline; }}
    @media (max-width:760px) {{
      h1 {{ font-size:36px; }}
      summary {{ grid-template-columns:1fr; }}
      .paper-card {{ grid-template-columns:1fr; }}
    }}
  </style>
</head>
<body>
  <header>
    <h1>Awesome BCI</h1>
    <p>A taxonomy-first map of Brain-Computer Interface research from {START_YEAR} through {END_YEAR}. Each year investigates up to {CANDIDATES_PER_YEAR} candidate papers, selects the top {TARGET_PER_YEAR}, and organizes the final papers by BCI research theme.</p>
    <nav>
      <a href="https://github.com/honggi82/awesome-BCI">README</a>
      <a href="data/{PAPERS_CSV}">CSV Dataset</a>
      <a href="data/{TAXONOMY_CSV}">Taxonomy CSV</a>
      <a href="data/{PERIOD_ANALYSIS_JSON}">Period Analysis JSON</a>
      <a href="#keywords-convention">Keywords Convention</a>
      <a href="data/{CANDIDATES_CSV}">Candidate Pool</a>
      <a href="paper/review_en.html">Review Paper</a>
      <a href="paper/review_ko.html">Korean Review</a>
    </nav>
  </header>
  <main>
    <div class="stats">
      <div class="stat"><strong id="statPapers">{len(flat)}</strong><span>selected papers</span></div>
      <div class="stat"><strong id="statYears">{len(YEARS)}</strong><span>years covered</span></div>
      <div class="stat"><strong id="statCitations">{total_cites:,}</strong><span>citation count total</span></div>
      <div class="stat"><strong id="statCategories">{len(cats)}</strong><span>topic categories</span></div>
    </div>
    <form class="filters" id="yearFilter">
      <div class="filter-field wide-field">
        <label for="periodPreset">Period</label>
        <select id="periodPreset" name="period">{period_options}</select>
      </div>
      <div class="filter-field">
        <label for="languageSelect">Language</label>
        <select id="languageSelect" name="lang">{language_options}</select>
      </div>
      <div class="filter-field">
        <label for="startYear">Start year</label>
        <select id="startYear" name="from">{start_year_options}</select>
      </div>
      <div class="filter-field">
        <label for="endYear">End year</label>
        <select id="endYear" name="to">{end_year_options}</select>
      </div>
      <button type="button" id="resetYears">Reset</button>
      <span id="rangeStatus"></span>
    </form>
    {research_overview}
    <section class="keyword-section" id="keywords-convention">
      <h2>Keywords Convention</h2>
      <p>These keyword tags follow the convention style used by AI-for-BCI awesome lists and define how papers can be labeled or scanned in this collection.</p>
      <div class="keyword-grid">{keyword_convention}</div>
      <p class="keyword-filter-status" id="keywordFilterStatus">Selected keyword: all | Matching papers: {len(flat):,} papers</p>
    </section>
    <h2>Taxonomy</h2>
    <p id="taxonomyTotalSummary"><strong>Total selected papers:</strong> {len(flat):,} papers; <strong>Categories:</strong> {len(cats)} categories.</p>
    <p>Each taxonomy section lists papers with publication year, journal or venue, citation count, main idea, strengths, limitations, and paper links. Sections are collapsed by default to keep the page scannable.</p>
    <div class="figures">
      <figure class="chart-figure">
        <img id="categoryDistributionChart" src="{period_chart_src('category_distribution', START_YEAR, END_YEAR)}" alt="Category distribution chart for {YEAR_RANGE_TEXT}">
        <figcaption id="categoryChartCaption">Category distribution ({YEAR_RANGE_TEXT})</figcaption>
      </figure>
      <figure class="chart-figure">
        <img id="yearlyCitationsChart" src="{period_chart_src('yearly_citations', START_YEAR, END_YEAR)}" alt="Yearly citation chart for {YEAR_RANGE_TEXT}">
        <figcaption id="citationChartCaption">Yearly citation mass ({YEAR_RANGE_TEXT})</figcaption>
      </figure>
    </div>
    {''.join(sections)}
  </main>
{year_filter_script}
</body>
</html>
"""
    (DOCS_DIR / "index.html").write_text(html_doc, encoding="utf-8")
    (DOCS_DIR / ".nojekyll").write_text("", encoding="utf-8")


def write_charts(flat):
    import matplotlib

    matplotlib.use("Agg")
    import matplotlib.pyplot as plt

    assets = DOCS_DIR / "assets"
    assets.mkdir(exist_ok=True)
    period_assets = assets / "periods"
    period_assets.mkdir(exist_ok=True)
    for stale in period_assets.glob("category_distribution_*.png"):
        stale.unlink()
    for stale in period_assets.glob("yearly_citations_*.png"):
        stale.unlink()

    def render_category_chart(rows, start, end, target):
        cats = category_stats(rows).most_common()
        labels = [c for c, _ in cats][::-1] or ["No papers"]
        values = [v for _, v in cats][::-1] or [0]
        fig, ax = plt.subplots(figsize=(11, 6), dpi=150)
        ax.barh(labels, values, color="#0f766e")
        ax.set_xlabel("Selected papers")
        ax.set_title(f"BCI Paper Taxonomy, {period_label(start, end)}")
        ax.grid(axis="x", alpha=0.25)
        fig.tight_layout()
        fig.savefig(target)
        plt.close(fig)

    def render_citation_chart(rows, start, end, target):
        stats = year_stats(rows)
        years = list(range(start, end + 1))
        citations = [stats.get(year, {}).get("citations", 0) for year in years]
        fig, ax = plt.subplots(figsize=(9, 5), dpi=150)
        ax.bar([str(year) for year in years], citations, color="#7c3aed")
        ax.set_ylabel("Citation count in selected set")
        ax.set_title(f"Yearly Citation Mass, {period_label(start, end)}")
        ax.grid(axis="y", alpha=0.25)
        if len(years) > 8:
            ax.tick_params(axis="x", labelrotation=45)
        fig.tight_layout()
        fig.savefig(target)
        plt.close(fig)

    for start, end in all_period_ranges():
        rows = [p for p in flat if start <= p["year"] <= end]
        category_target = period_assets / period_chart_filename("category_distribution", start, end)
        citation_target = period_assets / period_chart_filename("yearly_citations", start, end)
        render_category_chart(rows, start, end, category_target)
        render_citation_chart(rows, start, end, citation_target)

    shutil.copyfile(
        period_assets / period_chart_filename("category_distribution", START_YEAR, END_YEAR),
        assets / "category_distribution.png",
    )
    shutil.copyfile(
        period_assets / period_chart_filename("yearly_citations", START_YEAR, END_YEAR),
        assets / "yearly_citations.png",
    )


def reference_line(p):
    authors = p["authors"] or "Unknown authors"
    venue = p["venue"] or "Unknown venue"
    link = p["url"] or p["semanticScholarUrl"]
    return f"{authors}. ({p['year']}). {p['title']}. {venue}. {link}"


def review_sections(flat, korean=False):
    stats = year_stats(flat)
    cats = category_stats(flat)
    top_by_year = [stats[y]["top"] for y in YEARS if stats.get(y, {}).get("top")]
    if korean:
        title = f"{YEAR_RANGE_TEXT}년 뇌-컴퓨터 인터페이스 연구 동향: 공개 메타데이터 기반 리뷰"
        abstract = (
            f"이 리뷰 초안은 {START_YEAR}년부터 {END_YEAR}년까지의 Brain-Computer Interface(BCI) 연구를 "
            f"연도별 최대 {CANDIDATES_PER_YEAR}편의 후보로 조사하고 인용회수 기준으로 {TARGET_PER_YEAR}편씩 선별해 정리한다. 선별은 무료 공개 Semantic Scholar 메타데이터를 "
            "사용했으며, 검색어 기반 후보군을 연도 필터, 강화된 BCI 관련성 필터, 중복 제거, 후보 감사용 중요도 점수화로 처리했다. "
            "결과는 운동상상/운동 디코딩, SSVEP/P300/ERP, 재활과 신경보철, 침습형 인터페이스, 딥러닝, EEG 신호처리, "
            "언어/의사소통 BCI, 폐루프/하이브리드 BCI로 나뉜다."
        )
        methods = (
            "방법: 각 연도에 대해 BCI 관련 질의를 Semantic Scholar Academic Graph bulk search에 보내고, "
            "제목 또는 초록에 BCI 관련 핵심 표현이 있는 논문만 남겼다. DOI, arXiv, PubMed, CorpusId, paperId 순서로 "
            f"중복을 제거했다. 이후 연도별 최대 {CANDIDATES_PER_YEAR}편의 후보 풀을 만들고, 그 안에서 인용회수 상위 {TARGET_PER_YEAR}편을 최종 목록에 사용했다. "
            "영향력 있는 인용 수와 중요도 점수는 동률 처리와 감사용 보조 신호로 유지했다. 중요도 점수는 로그 변환 인용 수, 영향력 있는 인용 수, 주요 venue 신호, BCI 핵심어 밀도, 리뷰/서베이, 데이터셋/벤치마크, "
            f"임상/재활 관련성, 침습형 고대역폭 BCI, 최신 머신러닝 방법 보너스를 합산했다. {END_YEAR}년은 {date.today().isoformat()} 현재의 부분 연도라는 점에 유의해야 한다."
        )
        trends_intro = "주요 경향은 다음과 같다."
        caveat = (
            "한계: 이 문서는 전문 PDF를 모두 읽은 체계적 문헌고찰이 아니라 메타데이터 기반 리뷰 초안이다. "
            "인용 수는 최신 논문에 불리하며, 일부 논문의 venue/date 정보는 원천 데이터에서 누락될 수 있다."
        )
        conclusion = (
            f"결론적으로 {START_YEAR}년 이후 BCI 연구는 딥러닝 기반 EEG 디코딩과 벤치마크, 재활 응용, 침습형 고성능 디코딩, "
            "사용자 친화적 의사소통 시스템 사이에서 빠르게 확장되고 있다. 향후 리뷰에서는 실제 PDF 기반 품질 평가, "
            "임상 전환성, 데이터셋 재현성, 장기 안정성, 안전성과 윤리 문제를 함께 검토해야 한다."
        )
    else:
        title = f"Brain-Computer Interface Research from {START_YEAR} to {END_YEAR}: A Metadata-Driven Review"
        abstract = (
            f"This draft review maps Brain-Computer Interface (BCI) research from {START_YEAR} through {END_YEAR}, investigating up to "
            f"{CANDIDATES_PER_YEAR} candidate papers per year from free public Semantic Scholar metadata and selecting {TARGET_PER_YEAR} papers per year by citation count. Candidate papers were retrieved with BCI-related "
            "queries, filtered by target year and strengthened BCI relevance terms, deduplicated, and scored for candidate auditing. The resulting "
            "collection highlights work on motor imagery and movement decoding, SSVEP/P300/ERP systems, rehabilitation and "
            "neuroprosthetics, invasive interfaces, deep learning, EEG signal processing, speech and communication BCIs, and "
            "hybrid or closed-loop systems."
        )
        methods = (
            "Methods: For each year, BCI-oriented queries were sent to the Semantic Scholar Academic Graph bulk search endpoint. "
            "Records were retained when their title or abstract matched BCI relevance expressions, deduplicated by DOI, arXiv, "
            f"PubMed, CorpusId, or paperId, and reduced to at most {CANDIDATES_PER_YEAR} candidate papers per year. The final {TARGET_PER_YEAR} papers per year were selected by citation count, with influential citation count and the metadata importance score retained as tie-breakers and audit signals. The importance score combines log-scaled citations, log-scaled influential citations, recognized venue signals, BCI relevance-term density, and bonuses for reviews/surveys, datasets/benchmarks, clinical or rehabilitation relevance, invasive high-bandwidth BCIs, and modern machine-learning methods. The year {END_YEAR} should be interpreted as a partial year as of {date.today().isoformat()}."
        )
        trends_intro = "The main metadata-level trends are as follows."
        caveat = (
            "Limitations: this is a metadata-driven review draft rather than a full systematic review of every PDF. Citation counts "
            f"favor older work, recent {END_YEAR} papers are structurally under-counted, and some venues or publication dates are missing upstream."
        )
        conclusion = (
            f"Overall, BCI research since {START_YEAR} is expanding across deep-learning EEG decoding, benchmarks and datasets, rehabilitation, "
            "high-performance invasive decoding, and usable communication systems. A full manuscript should next add PDF-level appraisal, "
            "clinical translation evidence, dataset reproducibility, long-term stability, safety, and ethics."
        )
    category_lines = [f"{cat}: {count}" for cat, count in cats.most_common()]
    year_lines = [
        f"{y}: {stats[y]['count']} papers, {stats[y]['citations']:,} citations in selected set, top paper: {stats[y]['top']['title']}"
        for y in YEARS if y in stats
    ]
    refs = [reference_line(p) for p in top_by_year]
    return title, abstract, methods, trends_intro, category_lines, year_lines, caveat, conclusion, refs


def review_deep_dive(flat, korean=False):
    cats = category_stats(flat)
    stats = year_stats(flat)
    top_scored = sorted(flat, key=lambda p: p["importanceScore"], reverse=True)[:12]
    top_cited = sorted(flat, key=lambda p: p["citationCount"], reverse=True)[:12]
    leading_cat, leading_count = cats.most_common(1)[0]
    total_cites = sum(p["citationCount"] for p in flat)
    peak_year = max(stats, key=lambda y: stats[y]["citations"])
    newest_year = max(YEARS)

    if korean:
        findings = [
            f"선정된 {len(flat)}편은 총 {total_cites:,}회의 인용을 포함하며, 인용량은 {peak_year}년에 가장 높다.",
            f"가장 큰 축은 {leading_cat}({leading_count}편)으로, {START_YEAR}년 이후 BCI 연구가 여전히 운동 의도 해석과 운동 기능 보조를 중심으로 조직되어 있음을 보여준다.",
            f"{newest_year}년 논문은 부분 연도라 인용 수가 낮지만, foundation model, 장기 안정성, 고성능 침습형 디코딩, 임상 전환성 같은 주제가 두드러진다.",
            "비침습 EEG-BCI는 데이터셋/벤치마크와 딥러닝 방법론이 빠르게 늘었고, 침습형 BCI는 의사소통·운동 복원 성능을 실제 사용성 문제와 연결하는 방향으로 이동하고 있다.",
        ]
        category_discussion = [
            f"{cat}: {count}편. 이 범주는 최종 목록의 {count / len(flat):.1%}를 차지한다."
            for cat, count in cats.most_common()
        ]
        future = [
            "PDF 전문 기반의 임상근거 수준 평가와 risk-of-bias 코딩을 추가한다.",
            "데이터셋 공개 여부, 재현 코드, 피험자 수, 장기 안정성 지표를 별도 열로 정규화한다.",
            "후보 500편 풀에서 연도별 저인용 최신 논문의 잠재력을 전문가 검토로 보정한다.",
            "장애 당사자 중심 사용성, 안전성, 개인정보, 신경윤리 기준을 별도 taxonomy로 확장한다.",
        ]
        top_scored_heading = "메타데이터 중요도 점수 상위 논문(감사용)"
        top_cited_heading = "인용 수 상위 논문"
    else:
        findings = [
            f"The {len(flat)} selected papers account for {total_cites:,} citations in the selected set, with the largest citation mass in {peak_year}.",
            f"The dominant category is {leading_cat} ({leading_count} papers), indicating that movement intention decoding and motor assistance remain the organizing center of BCI research since {START_YEAR}.",
            f"Papers from {newest_year} are structurally citation-disadvantaged because the year is partial, but they surface emerging themes around foundation models, stability, invasive high-bandwidth decoding, and clinical translation.",
            "Non-invasive EEG-BCI work is increasingly shaped by datasets, benchmarks, and deep learning, while invasive BCI work is moving from peak decoding performance toward communication, autonomy, and usability constraints.",
        ]
        category_discussion = [
            f"{cat}: {count} papers, representing {count / len(flat):.1%} of the selected set."
            for cat, count in cats.most_common()
        ]
        future = [
            "Add PDF-level appraisal of clinical evidence and risk-of-bias indicators.",
            "Normalize dataset availability, released code, participant counts, and long-term stability metrics.",
            "Use expert review to compensate for low citation counts among very recent papers in the 500-candidate pools.",
            "Extend the taxonomy with user-centered usability, safety, privacy, and neuroethics criteria.",
        ]
        top_scored_heading = "Top Papers by Metadata Importance Score (Audit Signal)"
        top_cited_heading = "Top Papers by Citation Count"

    return {
        "findings": findings,
        "category_discussion": category_discussion,
        "future": future,
        "top_scored": top_scored,
        "top_cited": top_cited,
        "top_scored_heading": top_scored_heading,
        "top_cited_heading": top_cited_heading,
    }


def html_ranked_table(rows, metric):
    heading = "Importance" if metric == "importance" else "Citations"
    out = [
        "<table>",
        f"<thead><tr><th>Year</th><th>Rank</th><th>Paper</th><th>{heading}</th><th>Category</th></tr></thead>",
        "<tbody>",
    ]
    for p in rows:
        metric_value = p["importanceScore"] if metric == "importance" else p["citationCount"]
        link = f'<a href="{html.escape(p["url"])}">{html.escape(p["title"])}</a>' if p["url"] else html.escape(p["title"])
        out.append(
            f"<tr><td>{p['year']}</td><td>{p['rank']}</td><td>{link}</td>"
            f"<td>{metric_value}</td><td>{html.escape(p['category'])}</td></tr>"
        )
    out.extend(["</tbody>", "</table>"])
    return "\n".join(out)


def write_review_html(flat, korean=False):
    title, abstract, methods, trends_intro, category_lines, year_lines, caveat, conclusion, refs = review_sections(flat, korean)
    deep = review_deep_dive(flat, korean)
    lang = "ko" if korean else "en"
    heading_refs = "선정 참고문헌 예시" if korean else "Selected References"
    heading_findings = "핵심 발견" if korean else "Key Findings"
    heading_category = "분야별 해석" if korean else "Category-Level Interpretation"
    heading_future = "향후 연구 의제" if korean else "Future Research Agenda"
    html_doc = f"""<!doctype html>
<html lang="{lang}">
<head>
  <meta charset="utf-8">
  <meta name="viewport" content="width=device-width, initial-scale=1">
  <title>{html.escape(title)}</title>
  <style>
    body {{ font-family: Georgia, 'Noto Serif KR', serif; max-width: 920px; margin: 40px auto; padding: 0 22px; line-height: 1.72; color:#172033; }}
    h1 {{ line-height:1.15; }}
    h2 {{ margin-top: 34px; }}
    li {{ margin: 6px 0; }}
    .abstract {{ background:#f5f7fb; border-left:4px solid #0f766e; padding:14px 18px; }}
    table {{ width:100%; border-collapse:collapse; margin:16px 0; }}
    th,td {{ border-bottom:1px solid #d9dee8; padding:8px; vertical-align:top; text-align:left; }}
    th {{ background:#f4f6fa; }}
  </style>
</head>
<body>
  <h1>{html.escape(title)}</h1>
  <p><strong>Generated:</strong> {date.today().isoformat()} &middot; <strong>Dataset:</strong> {len(flat)} papers</p>
  <h2>Abstract</h2>
  <p class="abstract">{html.escape(abstract)}</p>
  <h2>1. Introduction and Scope</h2>
  <p>{html.escape(methods)}</p>
  <h2>2. {heading_findings}</h2>
  <ul>{''.join(f'<li>{html.escape(x)}</li>' for x in deep['findings'])}</ul>
  <h2>3. Taxonomy</h2>
  <ul>{''.join(f'<li>{html.escape(x)}</li>' for x in category_lines)}</ul>
  <h2>4. {heading_category}</h2>
  <ul>{''.join(f'<li>{html.escape(x)}</li>' for x in deep['category_discussion'])}</ul>
  <h2>5. Year-by-Year Trends</h2>
  <p>{html.escape(trends_intro)}</p>
  <ul>{''.join(f'<li>{html.escape(x)}</li>' for x in year_lines)}</ul>
  <h2>6. {deep['top_cited_heading']}</h2>
  {html_ranked_table(deep['top_cited'], 'citations')}
  <h2>7. {deep['top_scored_heading']}</h2>
  {html_ranked_table(deep['top_scored'], 'importance')}
  <h2>8. {heading_future}</h2>
  <ul>{''.join(f'<li>{html.escape(x)}</li>' for x in deep['future'])}</ul>
  <h2>9. Limitations</h2>
  <p>{html.escape(caveat)}</p>
  <h2>10. Conclusion</h2>
  <p>{html.escape(conclusion)}</p>
  <h2>{heading_refs}</h2>
  <ol>{''.join(f'<li>{html.escape(ref)}</li>' for ref in refs)}</ol>
</body>
</html>
"""
    name = "review_ko.html" if korean else "review_en.html"
    (PAPER_DIR / name).write_text(html_doc, encoding="utf-8")


def write_review_docx(flat):
    title, abstract, methods, trends_intro, category_lines, year_lines, caveat, conclusion, refs = review_sections(flat, korean=False)
    deep = review_deep_dive(flat, korean=False)
    doc = Document()
    doc.add_heading(title, level=0)
    doc.add_paragraph(f"Generated: {date.today().isoformat()} | Dataset: {len(flat)} papers")
    doc.add_heading("Abstract", level=1)
    doc.add_paragraph(abstract)
    doc.add_heading("1. Introduction and Scope", level=1)
    doc.add_paragraph(methods)
    doc.add_heading("2. Key Findings", level=1)
    for line in deep["findings"]:
        doc.add_paragraph(line, style="List Bullet")
    doc.add_heading("3. Taxonomy", level=1)
    for line in category_lines:
        doc.add_paragraph(line, style="List Bullet")
    doc.add_heading("4. Category-Level Interpretation", level=1)
    for line in deep["category_discussion"]:
        doc.add_paragraph(line, style="List Bullet")
    doc.add_heading("5. Year-by-Year Trends", level=1)
    doc.add_paragraph(trends_intro)
    for line in year_lines:
        doc.add_paragraph(line, style="List Bullet")
    doc.add_heading("6. Top Papers by Citation Count", level=1)
    for p in deep["top_cited"]:
        doc.add_paragraph(f"{p['year']} #{p['rank']}: {p['title']} ({p['citationCount']} citations)", style="List Number")
    doc.add_heading("7. Top Papers by Metadata Importance Score (Audit Signal)", level=1)
    for p in deep["top_scored"]:
        doc.add_paragraph(f"{p['year']} #{p['rank']}: {p['title']} ({p['importanceScore']})", style="List Number")
    doc.add_heading("8. Future Research Agenda", level=1)
    for line in deep["future"]:
        doc.add_paragraph(line, style="List Bullet")
    doc.add_heading("9. Limitations", level=1)
    doc.add_paragraph(caveat)
    doc.add_heading("10. Conclusion", level=1)
    doc.add_paragraph(conclusion)
    doc.add_heading("Selected References", level=1)
    for ref in refs:
        doc.add_paragraph(ref, style="List Number")
    doc.save(PAPER_DIR / "review_en.docx")


def main():
    for path in (DATA_DIR, DOCS_DIR, PAPER_DIR):
        path.mkdir(exist_ok=True)
    selected, candidates = collect_papers()
    flat = write_json_csv(selected, candidates)
    write_taxonomy_dataset(flat)
    write_period_analysis(flat)
    write_readme(flat)
    write_charts(flat)
    write_site(flat)
    write_review_html(flat, korean=False)
    write_review_html(flat, korean=True)
    write_review_docx(flat)
    shutil.copyfile(DATA_DIR / PAPERS_CSV, DOCS_DIR / "data" / PAPERS_CSV)
    shutil.copyfile(DATA_DIR / TAXONOMY_CSV, DOCS_DIR / "data" / TAXONOMY_CSV)
    shutil.copyfile(DATA_DIR / CANDIDATES_CSV, DOCS_DIR / "data" / CANDIDATES_CSV)
    shutil.copyfile(PAPER_DIR / "review_en.html", DOCS_DIR / "paper" / "review_en.html")
    shutil.copyfile(PAPER_DIR / "review_ko.html", DOCS_DIR / "paper" / "review_ko.html")
    (ROOT / "LICENSE").write_text("CC-BY-4.0 for text and metadata curation; upstream paper metadata belongs to original sources.\n", encoding="utf-8")
    print(f"[done] generated {len(flat)} selected papers", flush=True)


if __name__ == "__main__":
    main()
